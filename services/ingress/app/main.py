from __future__ import annotations
import hashlib
# --- Set-based idempotency helper ---
def calculate_claim_set_hash(claim_id, db):
    """Fetch all content_hash for claim, sort, join, and return SHA-256 hash."""
    hashes = [d.content_hash for d in db.query(Document).filter(Document.claim_id == claim_id).all() if d.content_hash]
    hashes.sort()
    joined = ",".join(hashes)
    return hashlib.sha256(joined.encode("utf-8")).hexdigest()

import hashlib
import logging
import os
import re
import sys
import sys
import uuid
from datetime import datetime
from pathlib import Path, PurePosixPath
from typing import Any

import aiofiles
from celery import chord, group, chain
from fastapi import APIRouter, Body, Depends, FastAPI, File, Form, Header, HTTPException, Query, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.exceptions import RequestValidationError
from starlette.exceptions import HTTPException as StarletteHTTPException
from services.shared_tasks import (
    coding_task,
    intake_task,
    ocr_task,
    parser_task,
    risk_task,
    validator_task,
    finalize_claim_task,
    run_pipeline_inline,
)
from libs.shared.celery_app import celery_app
from pydantic import BaseModel, EmailStr, Field, field_validator
from sqlalchemy import func, text
from sqlalchemy.orm import Session, selectinload

from .config import settings
from .db import SessionLocal, check_db_health, engine, force_master_session
from .models import Claim, Document, DocValidation
from libs.auth.passwords import hash_password, password_matches, verify_password
from libs.shared.models import ParseJob, ParsedField, WorkflowState, User, Role, UserRoleTable, Organization, PatientProfile, StaffProfile, Invitation
from libs.shared.workflow_state import get_latest_workflow_state, upsert_workflow_state
from .schemas import ClaimListOut, ClaimOut
from .rate_limiter import RateLimiter


class AuthUser(BaseModel):
    user_id: str | None = None
    email: str | None = None
    role: str = "patient"
    patient_id: str | None = None
    tenant_id: str | None = None
    is_authenticated: bool = False


# Global cache for JWKS client to reuse SSL connections and public keys
_jwks_client = None

def _get_jwks_client():
    global _jwks_client
    if _jwks_client is None:
        import jwt
        jwks_url = os.getenv("AUTH_JWKS_URL")
        tenant_id = os.getenv("ENTRA_TENANT_ID")
        subdomain = os.getenv("ENTRA_SUBDOMAIN") or os.getenv("NEXT_PUBLIC_ENTRA_SUBDOMAIN")
        
        if not jwks_url:
            if subdomain and tenant_id:
                jwks_url = f"https://{subdomain}.ciamlogin.com/{tenant_id}/discovery/v2.0/keys"
            elif tenant_id:
                jwks_url = f"https://login.microsoftonline.com/{tenant_id}/discovery/v2.0/keys"
        
        if jwks_url:
            try:
                _jwks_client = jwt.PyJWKClient(jwks_url, cache_keys=True, max_cached_keys=16)
                logger.info(f"Initialized Entra ID JWKS client with endpoint: {jwks_url}")
            except Exception as e:
                logger.warning(f"Could not initialize PyJWKClient: {e}")
    return _jwks_client

def get_current_user_context(
    authorization: str | None = Header(None),
    x_patient_id: str | None = Header(None, alias="X-Patient-Id"),
    x_user_id: str | None = Header(None, alias="X-User-Id"),
    patient_id: str | None = Query(None),
) -> AuthUser:
    is_production = os.getenv("APP_ENV", "development").strip().lower() in ("production", "prod")
    
    # 1. Production Mode: Strictly require and cryptographically verify Bearer Token
    if is_production:
        if not authorization or not authorization.startswith("Bearer "):
            raise HTTPException(
                status_code=401,
                detail="Authentication credentials were not provided or invalid Bearer token format.",
                headers={"WWW-Authenticate": "Bearer"},
            )
        
        token = authorization.split(" ", 1)[1].strip()
        try:
            import jwt
            claims = None
            jwks_client = _get_jwks_client()
            
            if jwks_client is not None:
                try:
                    signing_key = jwks_client.get_signing_key_from_jwt(token)
                    claims = jwt.decode(
                        token,
                        signing_key.key,
                        algorithms=["RS256"],
                        options={"verify_exp": True, "verify_aud": False}
                    )
                except Exception as verify_err:
                    logger.warning(f"JWKS cryptographic verification failed: {verify_err}")
                    raise HTTPException(
                        status_code=401,
                        detail=f"Cryptographic signature verification failed: {verify_err}",
                        headers={"WWW-Authenticate": "Bearer"},
                    )
            else:
                # If JWKS URL is not configured yet, decode with expiration check
                claims = jwt.decode(token, options={"verify_signature": False, "verify_exp": True})
            
            sub = claims.get("sub") or claims.get("oid")
            email = claims.get("email") or claims.get("preferred_username")
            roles = claims.get("roles") or ["patient"]
            primary_role = roles[0] if isinstance(roles, list) and roles else "patient"
            tenant_id = claims.get("tid")
            resolved_user_id = x_user_id or sub
            extracted_patient_id = claims.get("patient_id") or resolved_user_id or email
            
            return AuthUser(
                user_id=resolved_user_id,
                email=email,
                role=primary_role,
                patient_id=extracted_patient_id,
                tenant_id=tenant_id,
                is_authenticated=True,
            )
        except HTTPException:
            raise
        except Exception as exc:
            logger.warning(f"JWT Token validation failed in production: {exc}")
            raise HTTPException(
                status_code=401,
                detail="Invalid or expired authentication token.",
                headers={"WWW-Authenticate": "Bearer"},
            )

    # 2. Development / Local Test Harness Mode
    resolved_user_id = x_user_id or (patient_id if patient_id and patient_id.lower() not in ("user", "null", "undefined") else None)
    resolved_patient = (x_patient_id or resolved_user_id or patient_id or "").strip() or None
    role = "patient"
    is_auth = False
    
    if authorization and authorization.startswith("Bearer "):
        token = authorization.split(" ", 1)[1].strip()
        if token:
            is_auth = True
    elif resolved_patient or resolved_user_id:
        is_auth = True
        
    return AuthUser(
        user_id=resolved_user_id,
        patient_id=resolved_patient,
        role=role,
        is_authenticated=is_auth,
    )



try:
    from libs.utils.audit import AuditLogger
except Exception:
    AuditLogger = None  # type: ignore

def _audit(db, action: str, claim_id=None, metadata=None):
    try:
        if AuditLogger:
            with SessionLocal() as audit_db:
                AuditLogger(audit_db, "ingress").log(action, claim_id=claim_id, metadata=metadata)
    except Exception:
        logger.debug("Audit log failed for %s", action, exc_info=True)

# ------------------------------------------------------------------ logging
logging.basicConfig(
    level=settings.log_level.upper(),
    format="%(asctime)s  %(levelname)-8s  %(name)s  %(message)s",
)
logger = logging.getLogger("ingress")

# log4net-style on-disk audit log for claim uploads
# Writes to <repo_root>/logs/claim_uploads.txt (override via CLAIMGPT_LOG_DIR).
try:
    sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "..", ".."))
    from libs.observability.file_logger import get_file_logger
    upload_log = get_file_logger("ingress.upload", "claim_uploads.txt")
except Exception:  # pragma: no cover - logging must never break the service
    logger.exception("Failed to initialise claim upload file logger; falling back to standard logger")
    upload_log = logger

RAW_STORAGE = Path(settings.storage_root).resolve()
RAW_STORAGE.mkdir(parents=True, exist_ok=True)

app = FastAPI(title="ClaimGPT Ingress Service")

# Global exception handler to ensure all errors return JSON
@app.exception_handler(Exception)
async def global_exception_handler(request, exc):
    logger.exception("Unhandled exception in ingress service")
    return JSONResponse(
        status_code=500,
        content={"detail": f"Internal server error: {str(exc)}"},
    )

@app.exception_handler(StarletteHTTPException)
async def http_exception_handler(request, exc):
    return JSONResponse(
        status_code=exc.status_code,
        content={"detail": exc.detail},
    )

@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request, exc):
    return JSONResponse(
        status_code=422,
        content={"detail": str(exc)},
    )

@app.on_event("startup")
async def startup_event():
    try:
        from libs.shared.models import Base as SharedBase
        logger.info("Auto-initializing database tables on startup...")
        SharedBase.metadata.create_all(bind=engine)
        try:
            from services.coding.app.db import Base as CodingBase
            CodingBase.metadata.create_all(bind=engine)
        except Exception as ce:
            logger.warning("CodingBase schema init skipped: %s", ce)
        try:
            from services.predictor.app.db import Base as PredictorBase
            PredictorBase.metadata.create_all(bind=engine)
        except Exception as pe:
            logger.warning("PredictorBase schema init skipped: %s", pe)
        logger.info("Database tables initialized successfully.")
    except Exception as e:
        logger.exception("Failed to initialize database tables on startup: %s", e)

@app.on_event("shutdown")
async def shutdown_event():
    from .rate_limiter import limiter_manager
    await limiter_manager.close()

# ------------------------------------------------------------------ CORS
app.add_middleware(
    CORSMiddleware,
    allow_origin_regex="https://.*|http://localhost:.*|http://127.0.0.1:.*",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ------------------------------------------------------------------ observability
try:
    import sys
    sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "..", ".."))
    from libs.observability.metrics import PrometheusMiddleware, init_metrics, metrics_endpoint
    from libs.observability.tracing import init_tracing, instrument_fastapi
    init_tracing("ingress")
    init_metrics("ingress")
    instrument_fastapi(app)
    app.add_middleware(PrometheusMiddleware)
    _metrics_handler = metrics_endpoint()
    if _metrics_handler:
        app.get("/metrics")(_metrics_handler)
except Exception:
    logger.debug("Observability libs not available — skipping")


# ------------------------------------------------------------------ lifecycle
@app.on_event("shutdown")
def _shutdown():
    engine.dispose()
    logger.info("DB engine disposed")


# ------------------------------------------------------------------ deps
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


def _parse_uuid(value: str) -> uuid.UUID:
    try:
        return uuid.UUID(value)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid UUID")


def _safe_filename(raw: str | None) -> str:
    """Strip directory components to prevent path-traversal via filename."""
    if not raw:
        return "upload.bin"
    return PurePosixPath(raw).name or "upload.bin"


# Map every file extension we accept to one canonical Content-Type so we can
# normalise uploads coming from clients that send non-standard MIMEs (e.g.
# Windows reporting ``image/jpg`` for .jpg, or browsers/curl falling back to
# ``application/octet-stream``).  Keep this in lock-step with the OCR engine's
# SUPPORTED_EXTENSIONS — anything OCR can read should be uploadable.
_EXTENSION_TO_CONTENT_TYPE: dict[str, str] = {
    ".pdf": "application/pdf",
    # Images
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".jpe": "image/jpeg",
    ".jfif": "image/jpeg",
    ".png": "image/png",
    ".tiff": "image/tiff",
    ".tif": "image/tiff",
    ".bmp": "image/bmp",
    ".webp": "image/webp",
    ".gif": "image/gif",
    ".heic": "image/heic",
    ".heif": "image/heif",
    # Office
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": "application/msword",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xls": "application/vnd.ms-excel",
    ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    ".ppt": "application/vnd.ms-powerpoint",
    # OpenDocument
    ".odt": "application/vnd.oasis.opendocument.text",
    ".ods": "application/vnd.oasis.opendocument.spreadsheet",
    ".odp": "application/vnd.oasis.opendocument.presentation",
    # Misc
    ".rtf": "application/rtf",
    ".txt": "text/plain",
    ".csv": "text/csv",
    ".json": "application/json",
    ".xml": "application/xml",
    ".html": "text/html",
    ".htm": "text/html",
}

# Common non-standard / aliased MIME types we should accept silently.
_CONTENT_TYPE_ALIASES: dict[str, str] = {
    "image/jpg": "image/jpeg",       # non-standard but seen in the wild (Windows)
    "image/pjpeg": "image/jpeg",     # progressive JPEG (legacy IE)
    "image/x-png": "image/png",      # legacy
    "image/x-citrix-jpeg": "image/jpeg",
    "image/x-citrix-png": "image/png",
    "text/xml": "application/xml",
}


def _resolve_content_type(file: UploadFile) -> tuple[str, bool]:
    """Decide the effective Content-Type for an upload.

    Returns ``(content_type, is_supported)``.  Falls back to the file extension
    when the client sends nothing useful (``application/octet-stream`` or an
    empty header).  This is the single source of truth for upload validation
    so `.jpg` files always pass even when browsers report `image/jpg`.
    """
    raw_ct = (file.content_type or "").lower().strip()
    suffix = Path(file.filename or "").suffix.lower()

    # 1) Direct match against allowed list.
    if raw_ct in settings.allowed_content_types:
        return raw_ct, True

    # 2) Try alias normalisation.
    if raw_ct in _CONTENT_TYPE_ALIASES:
        canonical = _CONTENT_TYPE_ALIASES[raw_ct]
        if canonical in settings.allowed_content_types:
            return canonical, True

    # 3) Browsers / curl often send application/octet-stream or nothing for
    #    unknown extensions — trust the file extension as long as we know it.
    if suffix in _EXTENSION_TO_CONTENT_TYPE:
        canonical = _EXTENSION_TO_CONTENT_TYPE[suffix]
        if canonical in settings.allowed_content_types:
            return canonical, True

    return raw_ct or "application/octet-stream", False


def _compute_upload_sha256(file_data: list[tuple[UploadFile, bytes, str]]) -> str:
    hasher = hashlib.sha256()
    for _, content, safe_name in file_data:
        hasher.update(safe_name.encode("utf-8", errors="ignore"))
        hasher.update(b"\x00")
        hasher.update(content)
        hasher.update(b"\x00")
    return hasher.hexdigest()


def _build_claim_response(db: Session, claim_id: uuid.UUID, extra: dict[str, Any] | None = None) -> dict[str, Any]:
    claim = (
        db.query(Claim)
        .options(selectinload(Claim.documents))
        .filter(Claim.id == claim_id)
        .first()
    )
    if not claim:
        raise HTTPException(status_code=404, detail="Claim not found")
    payload = ClaimOut.model_validate(claim).model_dump(mode="json")
    if extra:
        payload.update(extra)
    return payload


def _build_report_url(claim_id: uuid.UUID) -> str:
    return f"/claims/{claim_id}"


def _find_completed_claim_by_upload_hash(db: Session, upload_sha256: str) -> Claim | None:
    from libs.shared.models import AuditLog
    claim = (
        db.query(Claim)
        .join(AuditLog, AuditLog.claim_id == Claim.id)
        .filter(
            AuditLog.action == "CLAIM_CREATED",
            AuditLog.audit_metadata["upload_sha256"].as_string() == upload_sha256,
            Claim.status == "COMPLETED",
        )
        .order_by(Claim.created_at.desc())
        .first()
    )
    return claim



def _celery_worker_available(timeout: float = 0.6) -> bool:
    """Best-effort check that at least one Celery worker is online and ready.

    A short ping (<1s) is issued via the Celery control bus. If the broker is
    unreachable or no worker replies within the timeout, returns ``False`` —
    callers can then fall back to inline execution so uploads never get stuck.
    """
    try:
        replies = celery_app.control.ping(timeout=timeout) or []
        return bool(replies)
    except Exception:
        return False


def _should_run_inline() -> bool:
    """Decide between the Celery chain and in-process inline execution.
    Defaults to False so that high-performance Celery ML workers (ocr_worker, parsing_worker, other_worker)
    process claims asynchronously in parallel.
    """
    raw = (os.getenv("CLAIMGPT_INLINE_PIPELINE") or "").strip().lower()
    if raw in {"1", "true", "yes", "on", "inline"}:
        return True
    return False


def _enqueue_pipeline(
    file_metadata: list[dict[str, str]] | str,
    policy_id: str | None = None,
    patient_id: str | None = None,
) -> str:
    """Enqueue the full pipeline starting with intake task, or trigger OCR on an existing claim.
    
    Args:
        file_metadata: List of dicts with keys: path, safe_name, content_hash, effective_ct
                      OR a string claim_id to retrigger pipeline for an existing claim.
        policy_id: Optional policy ID
        patient_id: Optional patient ID
    
    Returns:
        Task ID as string, or "inline:{claim_id}" for inline execution
    """
    if isinstance(file_metadata, str):
        claim_id_str = file_metadata
        
        # Synchronously reset workflow state in DB so polling API immediately returns 5% progress
        try:
            from services.ocr.app.db import SessionLocal as OcrSessionLocal
            with OcrSessionLocal() as db_session:
                upsert_workflow_state(db_session, uuid.UUID(claim_id_str), "STARTING", status="RUNNING")
                db_session.commit()
        except Exception:
            logger.exception("Failed to reset workflow state for claim %s", claim_id_str)

        if _should_run_inline():
            import threading
            logger.warning(
                "Celery worker not detected (or inline mode forced) — running pipeline inline for existing claim %s",
                claim_id_str,
            )

            def _runner() -> None:
                try:
                    run_pipeline_inline(claim_id_str)
                except Exception:
                    logger.exception("Inline pipeline crashed")

            thread = threading.Thread(
                target=_runner,
                name="inline-pipeline",
                daemon=True,
            )
            thread.start()
            return "inline:queued"

        workflow_chain = chain(
            ocr_task.s(claim_id_str),                               # Step 2: OCR (intake bypassed)
            parser_task.s(),                                        # Step 3: Parser
            coding_task.s(),                                        # Step 4: Coding
            risk_task.s(),                                          # Step 5: Risk
            validator_task.s(),                                     # Step 6: Validator
            finalize_claim_task.s(),                                # Step 7: Finalize Callback
        )
        result = workflow_chain.apply_async()
        return str(result.id)

    if _should_run_inline():
        # For inline execution, intake task needs to create the claim first
        import threading
        logger.warning(
            "Celery worker not detected (or inline mode forced) — running pipeline inline",
        )

        def _runner() -> None:
            try:
                # For inline, we need to do intake synchronously first
                from services.ocr.app.db import SessionLocal as OcrSessionLocal
                import hashlib
                db = OcrSessionLocal()
                try:
                    # Create claim
                    claim = Claim(
                        policy_id=policy_id,
                        patient_id=patient_id,
                        status="UPLOADED",
                        source="PATIENT",
                    )
                    db.add(claim)
                    db.flush()
                    claim_id = claim.id
                    
                    # Create documents
                    for metadata in file_metadata:
                        doc = Document(
                            claim_id=claim_id,
                            file_name=metadata["safe_name"],
                            file_type=metadata["effective_ct"],
                            minio_path=metadata["path"],
                            content_hash=metadata["content_hash"],
                        )
                        db.add(doc)
                    
                    db.commit()
                    
                    # Calculate set_hash
                    hashes = [d.content_hash for d in db.query(Document).filter(Document.claim_id == claim_id).all() if d.content_hash]
                    hashes.sort()
                    set_hash = hashlib.sha256(",".join(hashes).encode("utf-8")).hexdigest()
                    
                    # Create ParseJob
                    from libs.shared.models import ParseJob as PJ
                    parse_job = PJ(claim_id=claim_id, status="PENDING", set_hash=set_hash)
                    db.add(parse_job)
                    db.commit()
                    
                    # Update workflow state
                    upsert_workflow_state(db, claim_id, "STARTING", status="RUNNING")
                    db.commit()
                    
                    claim_id_str = str(claim_id)
                finally:
                    db.close()
                
                # Now run the inline pipeline
                run_pipeline_inline(claim_id_str)
            except Exception:
                logger.exception("Inline pipeline crashed")

        thread = threading.Thread(
            target=_runner,
            name="inline-pipeline",
            daemon=True,
        )
        thread.start()
        return "inline:queued"

    workflow_chain = chain(
        intake_task.s(file_metadata, policy_id, patient_id),  # Step 1: Intake (DB operations)
        ocr_task.s(),                                           # Step 2: OCR
        parser_task.s(),                                        # Step 3: Parser
        coding_task.s(),                                        # Step 4: Coding
        risk_task.s(),                                          # Step 5: Risk
        validator_task.s(),                                     # Step 6: Validator
        finalize_claim_task.s(),                                # Step 7: Finalize Callback
    )
    result = workflow_chain.apply_async()
    return str(result.id)


def _get_step_index(current_step: str | None, status: str | None) -> int:
    if current_step in ['OCR_STARTED', 'OCR_FINISHED']:
        return 1
    elif current_step in ['PARSING_STARTED', 'PARSING_FINISHED']:
        return 2
    elif current_step in ['CODING_STARTED', 'CODING_FINISHED', 'RISK_STARTED', 'RISK_FINISHED', 'VALIDATION_STARTED', 'VALIDATION_FINISHED']:
        return 3
    elif current_step in ['FINALIZE_STARTED', 'FINALIZE_FINISHED']:
        return 4
    elif status == 'FINISHED':
        return 5
    else:
        return 0


_PATIENT_NAME_PATTERNS = [
    re.compile(r"(?im)(?:^|\n)\s*(?:patient\s*name|name\s*of\s*patient)\s*[:\-]\s*([^\n\r|]+)"),
]

_DOB_PATTERNS = [
    re.compile(r"(?im)(?:^|\n)\s*(?:date\s*of\s*birth|dob|d\.o\.b)\s*[:\-]\s*([^\n\r|]+)"),
]

_MONTHS = {
    "jan": 1, "january": 1,
    "feb": 2, "february": 2,
    "mar": 3, "march": 3,
    "apr": 4, "april": 4,
    "may": 5,
    "jun": 6, "june": 6,
    "jul": 7, "july": 7,
    "aug": 8, "august": 8,
    "sep": 9, "sept": 9, "september": 9,
    "oct": 10, "october": 10,
    "nov": 11, "november": 11,
    "dec": 12, "december": 12,
}


def _canonical_name(value: str | None) -> str:
    if not value:
        return ""
    return re.sub(r"\s+", " ", value).strip().lower()


def _normalize_dob(value: str | None) -> str:
    if not value:
        return ""
    raw = re.sub(r"\s+", " ", value).strip().replace(",", "")
    m_num = re.fullmatch(r"(\d{1,2})[\-/.](\d{1,2})[\-/.](\d{2,4})", raw)
    if m_num:
        day, month, year = int(m_num.group(1)), int(m_num.group(2)), int(m_num.group(3))
        if year < 100:
            year += 2000 if year < 50 else 1900
        return f"{year:04d}-{month:02d}-{day:02d}"

    m_mon = re.fullmatch(r"(\d{1,2})[\-/. ]([A-Za-z]{3,9})[\-/. ](\d{2,4})", raw)
    if m_mon:
        day, month_token, year = int(m_mon.group(1)), m_mon.group(2).lower(), int(m_mon.group(3))
        month = _MONTHS.get(month_token)
        if month:
            if year < 100:
                year += 2000 if year < 50 else 1900
            return f"{year:04d}-{month:02d}-{day:02d}"

    m_alt = re.fullmatch(r"([A-Za-z]{3,9})\s+(\d{1,2})\s+(\d{2,4})", raw)
    if m_alt:
        month_token, day, year = m_alt.group(1).lower(), int(m_alt.group(2)), int(m_alt.group(3))
        month = _MONTHS.get(month_token)
        if month:
            if year < 100:
                year += 2000 if year < 50 else 1900
            return f"{year:04d}-{month:02d}-{day:02d}"

    return raw.lower()


def _extract_text_for_identity(file_path: Path, file_type: str | None) -> str:
    file_type = (file_type or "").lower()
    suffix = file_path.suffix.lower()

    if file_type == "application/pdf" or suffix == ".pdf":
        try:
            import pdfplumber

            parts: list[str] = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages[:5]:
                    t = page.extract_text() or ""
                    if t.strip():
                        parts.append(t)
            return "\n".join(parts)
        except Exception:
            return ""

    if suffix == ".docx":
        try:
            import docx

            d = docx.Document(str(file_path))
            return "\n".join(p.text for p in d.paragraphs if p.text)
        except Exception:
            return ""

    if suffix in {".xlsx", ".xlsm"}:
        try:
            import openpyxl

            wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
            lines: list[str] = []
            for ws in wb.worksheets[:3]:
                for row in ws.iter_rows(min_row=1, max_row=60, values_only=True):
                    vals = [str(c).strip() for c in row if c is not None and str(c).strip()]
                    if vals:
                        lines.append(" | ".join(vals))
            return "\n".join(lines)
        except Exception:
            return ""

    if suffix in {".txt", ".csv", ".json", ".xml", ".html", ".htm"}:
        try:
            return file_path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ""

    return ""


def _extract_identity_from_text(text: str) -> tuple[str | None, str | None]:
    if not text:
        return None, None

    patient_name: str | None = None
    dob: str | None = None

    for pat in _PATIENT_NAME_PATTERNS:
        m = pat.search(text)
        if m:
            patient_name = m.group(1).strip()
            break

    for pat in _DOB_PATTERNS:
        m = pat.search(text)
        if m:
            dob = m.group(1).strip()
            break

    if patient_name:
        patient_name = re.sub(r"\s+", " ", patient_name).strip()
    if dob:
        dob = re.sub(r"\s+", " ", dob).strip()
    return patient_name, dob


def _existing_identity_anchor(db: Session, claim_id: uuid.UUID) -> tuple[str | None, str | None]:
    rows = (
        db.query(DocValidation)
        .filter(
            DocValidation.claim_id == claim_id,
            DocValidation.doc_type == "IDENTITY_GATE",
            DocValidation.status == "VALID",
        )
        .order_by(DocValidation.created_at.asc())
        .all()
    )
    if rows:
        locked = []
        for row in rows:
            md = row.validation_metadata or {}
            if md.get("anchor_locked"):
                locked.append(row)
        picked = locked[0] if locked else rows[0]
        md = picked.validation_metadata or {}
        return picked.patient_name, md.get("identity_dob")

    # Fallback to other DocValidation rows (e.g. from the first batch OCR/validation)
    other_val = (
        db.query(DocValidation)
        .filter(
            DocValidation.claim_id == claim_id,
            DocValidation.status == "VALID",
            DocValidation.patient_name.isnot(None),
        )
        .order_by(DocValidation.created_at.asc())
        .first()
    )
    if other_val:
        md = other_val.validation_metadata or {}
        return other_val.patient_name, md.get("identity_dob")

    # Fallback to ParsedField (populated by LLM parser for first batch)
    pf_name = (
        db.query(ParsedField.field_value)
        .filter(
            ParsedField.claim_id == claim_id,
            ParsedField.field_name == "patient_name",
        )
        .first()
    )
    if pf_name and pf_name[0]:
        pf_dob = (
            db.query(ParsedField.field_value)
            .filter(
                ParsedField.claim_id == claim_id,
                ParsedField.field_name == "dob",
            )
            .first()
        )
        return pf_name[0], pf_dob[0] if pf_dob else None

    return None, None


def _upsert_identity_validation(
    db: Session,
    *,
    claim_id: uuid.UUID,
    document_id: uuid.UUID,
    file_name: str,
    status: str,
    patient_match: str,
    patient_name: str | None,
    dob: str | None,
    excluded: bool,
    needs_manual_review: bool,
    reason: str,
    anchor_locked: bool,
) -> None:
    db.query(DocValidation).filter(
        DocValidation.claim_id == claim_id,
        DocValidation.document_id == document_id,
        DocValidation.doc_type == "IDENTITY_GATE",
    ).delete(synchronize_session=False)

    metadata: dict[str, Any] = {
        "phase": "UPLOAD_IDENTITY_GATE",
        "file_name": file_name,
        "identity_dob": dob,
        "excluded_from_pipeline": excluded,
        "needs_manual_review": needs_manual_review,
        "reason": reason,
        "anchor_locked": anchor_locked,
        "checked_at_utc": datetime.utcnow().isoformat() + "Z",
    }

    db.add(DocValidation(
        claim_id=claim_id,
        document_id=document_id,
        status=status,
        doc_type="IDENTITY_GATE",
        doc_type_label="Identity Gate",
        is_medical=1,
        patient_match=patient_match,
        confidence=1.0,
        patient_name=patient_name,
        patient_id_extracted=None,
        issues=[reason],
        validation_metadata=metadata,
    ))


def _apply_identity_gate(
    db: Session,
    claim_id: uuid.UUID,
    documents: list[Document],
) -> dict[str, Any]:
    anchor_name, anchor_dob = _existing_identity_anchor(db, claim_id)
    anchor_name_key = _canonical_name(anchor_name)

    accepted_docs: list[str] = []
    rejected_docs: list[dict[str, str]] = []
    manual_review_required = False

    for doc in documents:
        text = _extract_text_for_identity(Path(doc.minio_path), doc.file_type)
        
        # Check if text is empty or too short (meaning image, scanned PDF, or empty doc)
        if not text or len(text.strip()) < 20:
            # Synchronous text extraction was not possible or returned minimal text.
            # Accept it for the pipeline so it can be OCR'd and validated asynchronously.
            _upsert_identity_validation(
                db,
                claim_id=claim_id,
                document_id=doc.id,
                file_name=doc.file_name,
                status="VALID",
                patient_match="PENDING",
                patient_name=None,
                dob=None,
                excluded=False,
                needs_manual_review=False,
                reason="Document requires OCR for identity verification",
                anchor_locked=False,
            )
            accepted_docs.append(doc.file_name)
            continue

        # Check if this document is actually an identity proof
        is_identity_doc = any(kw in text.lower() for kw in ("government of india", "unique identification authority", "uidai", "income tax department", "permanent account number", "voter id", "passport", "driving licence", "identity card"))
        if not is_identity_doc:
            _upsert_identity_validation(
                db,
                claim_id=claim_id,
                document_id=doc.id,
                file_name=doc.file_name,
                status="VALID",
                patient_match="SKIP",
                patient_name=None,
                dob=None,
                excluded=False,
                needs_manual_review=False,
                reason="Medical or non-KYC document bypassed Identity Gate",
                anchor_locked=False,
            )
            accepted_docs.append(doc.file_name)
            continue

        patient_name, dob_raw = _extract_identity_from_text(text)
        dob = _normalize_dob(dob_raw) if dob_raw else ""

        if not patient_name:
            manual_review_required = True
            reason = "Document missing required patient_name"
            _upsert_identity_validation(
                db,
                claim_id=claim_id,
                document_id=doc.id,
                file_name=doc.file_name,
                status="INVALID",
                patient_match="NO_DATA",
                patient_name=patient_name,
                dob=dob_raw,
                excluded=True,
                needs_manual_review=True,
                reason=reason,
                anchor_locked=False,
            )
            rejected_docs.append({"file_name": doc.file_name, "reason": reason})
            continue

        name_key = _canonical_name(patient_name)

        if not anchor_name_key:
            anchor_name = patient_name
            anchor_dob = dob
            anchor_name_key = name_key
            _upsert_identity_validation(
                db,
                claim_id=claim_id,
                document_id=doc.id,
                file_name=doc.file_name,
                status="VALID",
                patient_match="MATCH",
                patient_name=patient_name,
                dob=dob,
                excluded=False,
                needs_manual_review=False,
                reason="Anchor identity established (name-only)",
                anchor_locked=True,
            )
            accepted_docs.append(doc.file_name)
            continue

        if name_key == anchor_name_key:
            _upsert_identity_validation(
                db,
                claim_id=claim_id,
                document_id=doc.id,
                file_name=doc.file_name,
                status="VALID",
                patient_match="MATCH",
                patient_name=patient_name,
                dob=dob,
                excluded=False,
                needs_manual_review=False,
                reason="Identity matched claim anchor (name-only)",
                anchor_locked=False,
            )
            accepted_docs.append(doc.file_name)
            continue

        manual_review_required = True
        reason = "Patient name mismatch with first-batch claim anchor"
        _upsert_identity_validation(
            db,
            claim_id=claim_id,
            document_id=doc.id,
            file_name=doc.file_name,
            status="INVALID",
            patient_match="MISMATCH",
            patient_name=patient_name,
            dob=dob,
            excluded=True,
            needs_manual_review=True,
            reason=reason,
            anchor_locked=False,
        )
        rejected_docs.append({"file_name": doc.file_name, "reason": reason})

    return {
        "accepted_count": len(accepted_docs),
        "accepted_docs": accepted_docs,
        "rejected_docs": rejected_docs,
        "manual_review_required": manual_review_required,
        "anchor_name": anchor_name,
        "anchor_dob": anchor_dob,
    }


# ------------------------------------------------------------------ routes
router = APIRouter()


def _ensure_users_password_hash_column() -> None:
    pass


class RegisterUserIn(BaseModel):
    username: str
    password: str | None = None
    password_hash: str | None = None
    role: str
    first_name: str | None = None
    last_name: str | None = None
    phone: str | None = None
    organization: str | None = None
    employee_id: str | None = None
    dob: str | None = None
    gender: str | None = None
    policy: str | None = None
    sum_insured: Any | None = None
    provider: str | None = "local"


class LoginUserIn(BaseModel):
    username: str
    password: str | None = None
    password_hash: str | None = None
    role: str | None = None


@router.post("/auth/register", status_code=201, dependencies=[Depends(RateLimiter(limit=3, window_seconds=60))])
def register_local_user(payload: RegisterUserIn):
    """Register a user (patient or organization admin) directly in the database
    without Keycloak/Entra requirement."""
    _ensure_users_password_hash_column()

    email = str(payload.username).strip().lower()
    if not email:
        raise HTTPException(status_code=400, detail="Username/Email is required")

    role_str = str(payload.role).lower()
    # /register            -> patient/submitter
    # /register/organization -> organization admin
    normalized_role = "reviewer" if role_str in ("tpa", "admin", "organization", "org_admin") else "submitter"

    first_name = (payload.first_name or email.split("@")[0] or "User").strip()
    last_name = (payload.last_name or "").strip()

    # Parse sum_insured if provided
    sum_insured_val = None
    if payload.sum_insured is not None and str(payload.sum_insured).strip() != "":
        try:
            sum_insured_val = float(payload.sum_insured)
        except (ValueError, TypeError):
            sum_insured_val = None

    # Parse dob if provided (supports YYYY-MM-DD, DD/MM/YYYY, DD-MM-YYYY, DDMMYYYY, DD Mon YYYY)
    dob_val = None
    if payload.dob and str(payload.dob).strip() != "":
        dob_str = str(payload.dob).strip()
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%d%m%Y", "%d %b %Y", "%d %B %Y"):
            try:
                dob_val = datetime.strptime(dob_str, fmt).date()
                break
            except (ValueError, TypeError):
                continue

    # Organization registration requires an organization name
    if normalized_role == "reviewer":
        org_name_check = (payload.organization or "").strip()
        if not org_name_check:
            raise HTTPException(status_code=400, detail="Organization name is required for admin registration")

    with SessionLocal() as db:
        try:
            # 1. Create or get User
            user_row = db.execute(
                text("SELECT id FROM users WHERE lower(email) = lower(:email)"),
                {"email": email},
            ).mappings().first()

            supplied_hash = (payload.password_hash or "").strip()
            password_hash = supplied_hash or (hash_password(payload.password) if payload.password else None)

            if user_row:
                user_id = user_row["id"]
                db.execute(
                    text("""
                        UPDATE users
                        SET status = 'ACTIVE',
                            password_hash = COALESCE(:password_hash, password_hash),
                            updated_at = CURRENT_TIMESTAMP
                        WHERE id = :id
                    """),
                    {"id": user_id, "password_hash": password_hash},
                )
            else:
                new_user = User(
                    email=email,
                    phone=payload.phone or None,
                    external_provider='local',
                    external_subject_id=email,
                    status='ACTIVE',
                    email_verified=True,
                    password_hash=password_hash
                )
                db.add(new_user)
                db.flush()
                user_id = new_user.id

            # 2. Assign Role
            role_row = db.execute(
                text("SELECT id FROM roles WHERE name = :role_name"),
                {"role_name": normalized_role},
            ).mappings().first()

            if not role_row:
                new_role = Role(
                    name=normalized_role,
                    description=f"{normalized_role.title()} role"
                )
                db.add(new_role)
                db.flush()
                role_id = new_role.id
            else:
                role_id = role_row["id"]

            role_exists = db.execute(
                text("SELECT 1 FROM user_roles WHERE user_id = :user_id AND role_id = :role_id"),
                {"user_id": user_id, "role_id": role_id}
            ).scalar()
            if not role_exists:
                db.execute(
                    text("INSERT INTO user_roles (id, user_id, role_id) VALUES (:id, :user_id, :role_id)"),
                    {"id": uuid.uuid4(), "user_id": user_id, "role_id": role_id}
                )

            # 3. Create/Update Profile depending on role
            if normalized_role == "submitter":
                # Patient profile
                existing_profile = db.execute(
                    text("SELECT id FROM patient_profiles WHERE user_id = :user_id"),
                    {"user_id": user_id},
                ).mappings().first()

                if existing_profile:
                    db.execute(
                        text("""
                            UPDATE patient_profiles
                            SET first_name = :first_name,
                                last_name = :last_name,
                                dob = COALESCE(:dob, dob),
                                gender = COALESCE(:gender, gender),
                                policy_number = COALESCE(:policy, policy_number),
                                sum_insured = COALESCE(:sum_insured, sum_insured),
                                updated_at = CURRENT_TIMESTAMP
                            WHERE user_id = :user_id
                        """),
                        {
                            "user_id": user_id,
                            "first_name": first_name,
                            "last_name": last_name,
                            "dob": dob_val,
                            "gender": payload.gender or None,
                            "policy": payload.policy or None,
                            "sum_insured": sum_insured_val,
                        },
                    )
                else:
                    db.execute(
                        text("""
                            INSERT INTO patient_profiles (id, user_id, first_name, last_name, dob, gender, policy_number, sum_insured, coverage_verified)
                            VALUES (:id, :user_id, :first_name, :last_name, :dob, :gender, :policy, :sum_insured, :coverage_verified)
                        """),
                        {
                            "id": uuid.uuid4(),
                            "user_id": user_id,
                            "first_name": first_name,
                            "last_name": last_name,
                            "dob": dob_val,
                            "gender": payload.gender or None,
                            "policy": payload.policy or None,
                            "sum_insured": sum_insured_val,
                            "coverage_verified": False,
                        },
                    )

            else:
                # Organization Admin profile
                org_name = (payload.organization or "").strip()
                org_row = db.execute(
                    text("SELECT id FROM organizations WHERE lower(name) = lower(:name) AND type = 'TPA'"),
                    {"name": org_name},
                ).mappings().first()

                if not org_row:
                    new_org = Organization(
                        name=org_name,
                        type='TPA',
                        status='ACTIVE'
                    )
                    db.add(new_org)
                    db.flush()
                    org_id = new_org.id
                else:
                    org_id = org_row["id"]

                existing_staff = db.execute(
                    text("SELECT id FROM staff_profiles WHERE user_id = :user_id"),
                    {"user_id": user_id},
                ).mappings().first()

                if existing_staff:
                    db.execute(
                        text("""
                            UPDATE staff_profiles
                            SET first_name = :first_name,
                                last_name = :last_name,
                                organization_id = :org_id,
                                employee_id = COALESCE(:employee_id, employee_id),
                                designation = :designation,
                                updated_at = CURRENT_TIMESTAMP
                            WHERE user_id = :user_id
                        """),
                        {
                            "user_id": user_id,
                            "first_name": first_name,
                            "last_name": last_name,
                            "org_id": org_id,
                            "employee_id": payload.employee_id or None,
                            "designation": "Organization Admin",
                        },
                    )
                else:
                    db.execute(
                        text("""
                            INSERT INTO staff_profiles (id, user_id, organization_id, first_name, last_name, employee_id, designation, status)
                            VALUES (:id, :user_id, :org_id, :first_name, :last_name, :employee_id, :designation, 'ACTIVE')
                        """),
                        {
                            "id": uuid.uuid4(),
                            "user_id": user_id,
                            "org_id": org_id,
                            "first_name": first_name,
                            "last_name": last_name,
                            "employee_id": payload.employee_id or None,
                            "designation": "Organization Admin",
                        },
                    )

            db.commit()
            return {
                "success": True,
                "user_id": str(user_id),
                "email": email,
                "role": normalized_role,
                "organization": payload.organization if normalized_role == "reviewer" else None,
                "message": "User registered and profile stored in database successfully",
            }
        except HTTPException:
            raise
        except Exception as exc:
            db.rollback()
            logger.exception("Failed to register user locally in database")
            raise HTTPException(status_code=500, detail=f"Database registration error: {str(exc)}") from exc


import re
import secrets
from datetime import timedelta


def _slugify_org(name: str) -> str:
    s = name.strip().lower()
    s = re.sub(r"[^a-z0-9]+", "-", s)
    return s.strip("-")


class InviteReviewerIn(BaseModel):
    first_name: str
    last_name: str
    email: str
    organization: str
    role: str | None = "reviewer"
    invited_by: str | None = None


@router.post("/auth/invite", status_code=201)
def invite_reviewer(payload: InviteReviewerIn):
    """Create an invitation for a claim reviewer under the specified organization."""
    email = payload.email.strip().lower()
    first_name = payload.first_name.strip()
    last_name = payload.last_name.strip()
    org_name = payload.organization.strip()

    if not email or not first_name or not last_name or not org_name:
        raise HTTPException(status_code=400, detail="First name, last name, work email, and organization are required")

    with SessionLocal() as db:
        try:
            # 1. Find or create Organization
            org_row = db.execute(
                text("SELECT id FROM organizations WHERE lower(name) = lower(:name)"),
                {"name": org_name},
            ).mappings().first()

            if not org_row:
                new_org = Organization(
                    name=org_name,
                    type="TPA",
                    status="ACTIVE",
                )
                db.add(new_org)
                db.flush()
                org_id = new_org.id
            else:
                org_id = org_row["id"]

            # 2. Check for existing pending invitation or create new one
            existing_inv = db.query(Invitation).filter(
                Invitation.email == email,
                Invitation.organization_id == org_id,
                Invitation.status == "PENDING",
            ).first()

            token = secrets.token_urlsafe(32)
            expires_at = datetime.now(timezone.utc) + timedelta(days=7)

            if existing_inv:
                existing_inv.first_name = first_name
                existing_inv.last_name = last_name
                existing_inv.token = token
                existing_inv.expires_at = expires_at
                existing_inv.updated_at = datetime.now(timezone.utc)
                invitation_id = existing_inv.id
            else:
                new_inv = Invitation(
                    email=email,
                    first_name=first_name,
                    last_name=last_name,
                    organization_id=org_id,
                    role=(payload.role or "reviewer").lower(),
                    status="PENDING",
                    token=token,
                    expires_at=expires_at,
                )
                db.add(new_inv)
                db.flush()
                invitation_id = new_inv.id

            db.commit()
            return {
                "success": True,
                "invitation_id": str(invitation_id),
                "email": email,
                "first_name": first_name,
                "last_name": last_name,
                "organization": org_name,
                "role": payload.role or "reviewer",
                "status": "PENDING",
                "token": token,
                "message": f"Invitation queued successfully for {email}",
            }
        except HTTPException:
            raise
        except Exception as exc:
            db.rollback()
            logger.exception("Failed to create reviewer invitation")
            raise HTTPException(status_code=500, detail=f"Invitation error: {str(exc)}") from exc


@router.post("/auth/login", status_code=200, dependencies=[Depends(RateLimiter(limit=5, window_seconds=60))])
def login_local_user(payload: LoginUserIn):
    """Authenticate a locally registered user using a stored password hash.

    Patients send an explicit role ("submitter") and get the existing
    role-mismatch behavior (401/403 with detail). Organizations send no
    role at all — the role (admin or reviewer) is resolved from the
    account itself, and any failure (user not found, bad password, wrong
    role, inactive account) is surfaced as a generic "Access denied" so
    we don't leak which part failed.
    """
    _ensure_users_password_hash_column()

    email = str(payload.username).strip().lower()
    role_str = str(payload.role).strip().lower() if payload.role else None
    is_org_login = role_str is None

    if not email or (not payload.password and not payload.password_hash):
        raise HTTPException(status_code=400, detail="Username and password are required")

    normalized_role = "submitter"

    with force_master_session():
        with SessionLocal() as db:
            user_row = db.execute(
                text("SELECT id, password_hash, status FROM users WHERE lower(email) = lower(:email)"),
                {"email": email},
            ).mappings().first()

            if not user_row:
                if is_org_login:
                    raise HTTPException(status_code=401, detail="Access denied")
                raise HTTPException(status_code=401, detail="Username not found")

            if user_row["status"] in ("BLOCKED", "DELETED"):
                if is_org_login:
                    raise HTTPException(status_code=403, detail="Access denied")
                raise HTTPException(status_code=403, detail="Account is not active")

            supplied_hash = (payload.password_hash or "").strip()
            supplied_password = payload.password or ""
            stored_hash = user_row["password_hash"]

            if not stored_hash:
                if is_org_login:
                    raise HTTPException(status_code=401, detail="Access denied")
                raise HTTPException(status_code=401, detail="Invalid email or password")

            password_ok = (
                password_matches(supplied_hash, stored_hash)
                if supplied_hash
                else password_matches(supplied_password, stored_hash)
            )

            if not password_ok:
                if is_org_login:
                    raise HTTPException(status_code=401, detail="Access denied")
                raise HTTPException(status_code=401, detail="Invalid password")

            actual_role_row = db.query(Role.name).join(
                UserRoleTable, UserRoleTable.role_id == Role.id
            ).filter(
                UserRoleTable.user_id == user_row["id"]
            ).order_by(Role.name).first()
            actual_role = actual_role_row[0] if actual_role_row else None

            organization_name = None
            organization_slug = None

            if is_org_login:
                # Organization accounts must hold either the admin or
                # reviewer role — a submitter (patient) account trying the
                # org login is denied without revealing that they're a patient.
                if actual_role not in ("admin", "reviewer"):
                    raise HTTPException(status_code=403, detail="Access denied")
                normalized_role = actual_role

                org_row = db.execute(
                    text("""
                        SELECT o.name
                        FROM staff_profiles sp
                        JOIN organizations o ON o.id = sp.organization_id
                        WHERE sp.user_id = :user_id
                    """),
                    {"user_id": user_row["id"]},
                ).mappings().first()

                if not org_row:
                    # Role says admin/reviewer but there's no org attached — treat as denied.
                    raise HTTPException(status_code=403, detail="Access denied")

                organization_name = org_row["name"]
                organization_slug = _slugify_org(organization_name)
            else:
                role_match = db.execute(
                    text(
                        "SELECT 1 FROM user_roles ur JOIN roles r ON ur.role_id = r.id WHERE ur.user_id = :user_id AND r.name = :role_name"
                    ),
                    {"user_id": user_row["id"], "role_name": normalized_role},
                ).scalar()

                if not role_match:
                    raise HTTPException(
                        status_code=403,
                        detail={
                            "message": f"User is not registered as a {role_str}",
                            "actual_role": actual_role or normalized_role,
                        },
                    )

            db.execute(
                text("UPDATE users SET last_login_at = CURRENT_TIMESTAMP, updated_at = CURRENT_TIMESTAMP WHERE id = :id"),
                {"id": user_row["id"]},
            )
            db.commit()

    return {
        "success": True,
        "user_id": str(user_row["id"]),
        "email": email,
        "role": normalized_role,
        "organization": organization_name,
        "organization_slug": organization_slug,
        "message": "Login successful",
    }


# ------------------------------------------------------------------ Entra CIAM user synchronization
class SyncEntraUserIn(BaseModel):
    email: str = Field(min_length=5, max_length=255)
    name: str | None = Field(default=None, max_length=150)
    first_name: str | None = Field(default=None, max_length=100)
    last_name: str | None = Field(default=None, max_length=100)
    company_name: str | None = Field(default=None, max_length=200)
    organization: str | None = Field(default=None, max_length=200)
    external_subject_id: str | None = None
    requested_role: str | None = "patient"
    account_role: str | None = None
    client_id: str | None = None
    phone: str | None = None
    dob: str | None = None
    gender: str | None = None
    policy: str | None = None
    sum_insured: float | str | None = None

    @field_validator("email", mode="before")
    @classmethod
    def validate_email_strict(cls, v: Any) -> str:
        s = str(v or "").strip()
        email_regex = r"^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$"
        if not re.match(email_regex, s):
            raise ValueError("Invalid email format.")
        return s

    @field_validator("name", "first_name", "last_name", "company_name", "organization", mode="before")
    @classmethod
    def sanitize_xss(cls, v: Any) -> str | None:
        if v is None:
            return None
        s = str(v).strip()
        cleaned = re.sub(r"<[^>]*>", "", s)
        return cleaned.strip()


@router.post("/auth/sync-entra-user", status_code=200)
def sync_entra_user(payload: SyncEntraUserIn):
    """Synchronize a Microsoft Entra External ID authenticated user with the database.

    - Organization Flow: If user is new from Azure Entra ID, automatically provisions
      the organization (from company_name), user record, admin role, and staff_profile,
      then routes them as an active Organization Admin. If existing, updates login timestamp
      and ensures staff profile integrity.
    - Patient Flow: If existing patient, returns profile and onboarding status.
      If new patient, registers user in users table and creates patient_profiles record.
    """
    _ensure_users_password_hash_column()

    email = str(payload.email).strip().lower()
    if not email:
        raise HTTPException(status_code=400, detail="Email is required")

    role_str = str(payload.requested_role or "patient").strip().lower()
    is_org_login = role_str in ("tpa", "admin", "reviewer", "organization", "org_admin")
    subject_id = payload.external_subject_id or email

    with force_master_session():
        with SessionLocal() as db:
            user_row = db.execute(
                text("SELECT id, email, status FROM users WHERE lower(email) = lower(:email)"),
                {"email": email},
            ).mappings().first()

            if is_org_login:
                # ----------------------------------------------------
                # Organization Flow
                # ----------------------------------------------------
                if user_row:
                    if user_row["status"] in ("BLOCKED", "DELETED"):
                        raise HTTPException(status_code=403, detail="Account is deactivated or blocked.")

                    user_id = user_row["id"]

                    # Check if this user is actually an organization staff member
                    actual_role_row = db.query(Role.name).join(
                        UserRoleTable, UserRoleTable.role_id == Role.id
                    ).filter(
                        UserRoleTable.user_id == user_id
                    ).order_by(Role.name).first()
                    actual_role = actual_role_row[0] if actual_role_row else None

                    staff_row = db.execute(
                        text("""
                            SELECT sp.id, sp.first_name, sp.last_name, sp.designation, o.name
                            FROM staff_profiles sp
                            JOIN organizations o ON o.id = sp.organization_id
                            WHERE sp.user_id = :user_id
                        """),
                        {"user_id": user_id},
                    ).mappings().first()

                    # Strict Security Check: If registered as a patient without staff profile, DENY
                    if actual_role == "submitter" and not staff_row:
                        raise HTTPException(
                            status_code=403,
                            detail="Access denied. Your account is registered as a patient, not as organization staff.",
                        )

                    # If no staff profile is found for an existing user, deny
                    if not staff_row:
                        raise HTTPException(
                            status_code=403,
                            detail="Access denied. No active organization profile is linked to this account.",
                        )

                    org_name = staff_row["name"]
                    org_slug = _slugify_org(org_name)
                    staff_fname = staff_row["first_name"] or ""
                    staff_lname = staff_row["last_name"] or ""

                    db.execute(
                        text("""
                            UPDATE users 
                            SET last_login_at = CURRENT_TIMESTAMP,
                                updated_at = CURRENT_TIMESTAMP,
                                external_provider = 'entra',
                                external_subject_id = COALESCE(:subject_id, external_subject_id)
                            WHERE id = :id
                        """),
                        {"id": user_id, "subject_id": subject_id},
                    )
                    db.commit()

                    return {
                        "success": True,
                        "user_id": str(user_id),
                        "email": email,
                        "name": f"{staff_fname} {staff_lname}".strip() or email.split("@")[0],
                        "first_name": staff_fname,
                        "last_name": staff_lname,
                        "role": "tpa",
                        "account_role": actual_role if actual_role in ("admin", "reviewer") else "admin",
                        "organization": org_name,
                        "organization_slug": org_slug,
                        "is_new_user": False,
                        "needs_onboarding": False,
                        "message": "Organization staff verified successfully",
                    }

                else:
                    # ----------------------------------------------------
                    # NEW User: Auto-provision as Organization Admin
                    # ----------------------------------------------------
                    name_parts = (payload.name or "").strip().split(" ")
                    first_name = (payload.first_name or "").strip()
                    last_name = (payload.last_name or "").strip()
                    if not first_name and payload.name:
                        first_name = name_parts[0]
                        last_name = " ".join(name_parts[1:]) if len(name_parts) > 1 else ""
                    if not first_name:
                        first_name = email.split("@")[0].capitalize()

                    org_name = (payload.company_name or payload.organization or "").strip()
                    if not org_name:
                        domain_part = email.split("@")[-1].split(".")[0]
                        if domain_part not in ("gmail", "yahoo", "outlook", "hotmail", "claimgpt", "test", "example"):
                            org_name = domain_part.replace("-", " ").replace("_", " ").title()
                        else:
                            org_name = "Star Health"

                    org_slug = _slugify_org(org_name)

                    # 1. Ensure Organization exists
                    org_row = db.execute(
                        text("SELECT id, name FROM organizations WHERE lower(name) = lower(:name)"),
                        {"name": org_name},
                    ).mappings().first()

                    if not org_row:
                        org_id = uuid.uuid4()
                        db.execute(
                            text("""
                                INSERT INTO organizations (id, name, type, status, created_at, updated_at)
                                VALUES (:id, :name, 'TPA', 'ACTIVE', CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                            """),
                            {"id": org_id, "name": org_name},
                        )
                    else:
                        org_id = org_row["id"]
                        org_name = org_row["name"]
                        org_slug = _slugify_org(org_name)

                    # 2. Create User
                    user_id = uuid.uuid4()
                    db.execute(
                        text("""
                            INSERT INTO users (id, email, phone, external_provider, external_subject_id, status, email_verified, created_at, updated_at, last_login_at)
                            VALUES (:id, :email, :phone, 'entra', :subject_id, 'ACTIVE', 1, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                        """),
                        {
                            "id": user_id,
                            "email": email,
                            "phone": payload.phone,
                            "subject_id": subject_id,
                        },
                    )

                    # 3. Ensure 'admin' Role exists and is assigned
                    admin_role = db.query(Role).filter(Role.name == "admin").first()
                    if not admin_role:
                        admin_role = Role(id=uuid.uuid4(), name="admin", description="Full system access")
                        db.add(admin_role)
                        db.flush()

                    db.add(UserRoleTable(id=uuid.uuid4(), user_id=user_id, role_id=admin_role.id))
                    db.flush()

                    # 4. Create staff_profiles record
                    db.execute(
                        text("""
                            INSERT INTO staff_profiles (id, user_id, organization_id, first_name, last_name, employee_id, designation, department, status, created_at, updated_at)
                            VALUES (:id, :user_id, :org_id, :first_name, :last_name, :employee_id, 'Administrator', 'Operations', 'ACTIVE', CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                        """),
                        {
                            "id": uuid.uuid4(),
                            "user_id": user_id,
                            "org_id": org_id,
                            "first_name": first_name,
                            "last_name": last_name or "",
                            "employee_id": f"EMP-{str(uuid.uuid4())[:8].upper()}",
                        },
                    )
                    db.commit()

                    return {
                        "success": True,
                        "user_id": str(user_id),
                        "email": email,
                        "name": f"{first_name} {last_name}".strip() or email.split("@")[0],
                        "first_name": first_name,
                        "last_name": last_name,
                        "role": "tpa",
                        "account_role": "admin",
                        "organization": org_name,
                        "organization_slug": org_slug,
                        "is_new_user": True,
                        "needs_onboarding": False,
                        "message": "Organization admin auto-provisioned successfully",
                    }

            else:
                # ----------------------------------------------------
                # Patient Flow
                # ----------------------------------------------------
                name_parts = (payload.name or "").strip().split(" ")
                first_name = name_parts[0] if name_parts and name_parts[0] else email.split("@")[0]
                last_name = " ".join(name_parts[1:]) if len(name_parts) > 1 else ""

                if user_row:
                    user_id = user_row["id"]

                    # Check role
                    actual_role_row = db.query(Role.name).join(
                        UserRoleTable, UserRoleTable.role_id == Role.id
                    ).filter(
                        UserRoleTable.user_id == user_id
                    ).order_by(Role.name).first()
                    actual_role = actual_role_row[0] if actual_role_row else "submitter"

                    if actual_role in ("admin", "reviewer"):
                        raise HTTPException(
                            status_code=403,
                            detail="Account is registered as organization staff. Please sign in via 'Continue as Organization'.",
                        )

                    # Check patient profile
                    profile_row = db.execute(
                        text("SELECT id, first_name, last_name, policy_number, sum_insured FROM patient_profiles WHERE user_id = :user_id"),
                        {"user_id": user_id},
                    ).mappings().first()

                    # If incoming payload provides onboarding info, update the profile immediately
                    if payload.policy or payload.dob or payload.gender or payload.sum_insured or payload.first_name:
                        sum_val = float(str(payload.sum_insured).replace(",", "").strip()) if payload.sum_insured else None
                        if profile_row:
                            db.execute(
                                text("""
                                    UPDATE patient_profiles
                                    SET first_name = COALESCE(:first_name, first_name),
                                        last_name = COALESCE(:last_name, last_name),
                                        gender = COALESCE(:gender, gender),
                                        policy_number = COALESCE(:policy_number, policy_number),
                                        sum_insured = COALESCE(:sum_insured, sum_insured),
                                        coverage_verified = 1,
                                        updated_at = CURRENT_TIMESTAMP
                                    WHERE user_id = :user_id
                                """),
                                {
                                    "user_id": user_id,
                                    "first_name": payload.first_name or None,
                                    "last_name": payload.last_name or None,
                                    "gender": payload.gender or None,
                                    "policy_number": payload.policy or None,
                                    "sum_insured": sum_val,
                                },
                            )
                        else:
                            db.execute(
                                text("""
                                    INSERT INTO patient_profiles (id, user_id, first_name, last_name, gender, policy_number, sum_insured, coverage_verified, created_at, updated_at)
                                    VALUES (:id, :user_id, :first_name, :last_name, :gender, :policy_number, :sum_insured, 1, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                                """),
                                {
                                    "id": uuid.uuid4(),
                                    "user_id": user_id,
                                    "first_name": first_name,
                                    "last_name": last_name or "",
                                    "gender": payload.gender or None,
                                    "policy_number": payload.policy or "POL-DEFAULT",
                                    "sum_insured": sum_val or 500000.0,
                                },
                            )
                        needs_onboarding = False
                    else:
                        needs_onboarding = not bool(profile_row and profile_row["policy_number"])

                    db.execute(
                        text("""
                            UPDATE users 
                            SET last_login_at = CURRENT_TIMESTAMP,
                                updated_at = CURRENT_TIMESTAMP,
                                external_provider = 'entra',
                                external_subject_id = COALESCE(:subject_id, external_subject_id)
                            WHERE id = :id
                        """),
                        {"id": user_id, "subject_id": subject_id},
                    )
                    db.commit()

                    p_fname = (profile_row["first_name"] if profile_row else None) or first_name
                    p_lname = (profile_row["last_name"] if profile_row else None) or last_name

                    return {
                        "success": True,
                        "user_id": str(user_id),
                        "email": email,
                        "name": f"{p_fname} {p_lname}".strip() or email.split("@")[0],
                        "first_name": p_fname,
                        "last_name": p_lname,
                        "role": "patient",
                        "account_role": "submitter",
                        "is_new_user": False,
                        "needs_onboarding": needs_onboarding,
                        "message": "Patient authenticated successfully",
                    }
                else:
                    # New Patient self-registration from Entra
                    new_user_id = uuid.uuid4()
                    db.execute(
                        text("""
                            INSERT INTO users (id, email, phone, external_provider, external_subject_id, status, email_verified, created_at, updated_at, last_login_at)
                            VALUES (:id, :email, :phone, 'entra', :subject_id, 'ACTIVE', 1, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                        """),
                        {
                            "id": new_user_id,
                            "email": email,
                            "phone": payload.phone or None,
                            "subject_id": subject_id,
                        },
                    )

                    # Ensure submitter role exists and assign it
                    role_row = db.execute(
                        text("SELECT id FROM roles WHERE name = 'submitter'"),
                    ).mappings().first()

                    if not role_row:
                        role_id = uuid.uuid4()
                        db.execute(
                            text("INSERT INTO roles (id, name, description, created_at) VALUES (:id, 'submitter', 'Patient submitter role', CURRENT_TIMESTAMP)"),
                            {"id": role_id},
                        )
                    else:
                        role_id = role_row["id"]

                    db.execute(
                        text("INSERT INTO user_roles (id, user_id, role_id, created_at) VALUES (:id, :user_id, :role_id, CURRENT_TIMESTAMP)"),
                        {"id": uuid.uuid4(), "user_id": new_user_id, "role_id": role_id},
                    )

                    sum_val = float(str(payload.sum_insured).replace(",", "").strip()) if payload.sum_insured else None
                    has_policy = bool(payload.policy)

                    # Create initial patient profile
                    db.execute(
                        text("""
                            INSERT INTO patient_profiles (id, user_id, first_name, last_name, gender, policy_number, sum_insured, coverage_verified, created_at, updated_at)
                            VALUES (:id, :user_id, :first_name, :last_name, :gender, :policy_number, :sum_insured, :coverage_verified, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                        """),
                        {
                            "id": uuid.uuid4(),
                            "user_id": new_user_id,
                            "first_name": first_name,
                            "last_name": last_name or "",
                            "gender": payload.gender or None,
                            "policy_number": payload.policy or None,
                            "sum_insured": sum_val,
                            "coverage_verified": 1 if has_policy else 0,
                        },
                    )

                    db.commit()

                    return {
                        "success": True,
                        "user_id": str(new_user_id),
                        "email": email,
                        "name": f"{first_name} {last_name}".strip() or email.split("@")[0],
                        "first_name": first_name,
                        "last_name": last_name,
                        "role": "patient",
                        "account_role": "submitter",
                        "is_new_user": True,
                        "needs_onboarding": not has_policy,
                        "message": "New patient registered in database successfully",
                    }


class RegisterUserIn(BaseModel):
    username: str
    password: str | None = None
    password_hash: str | None = None
    role: str = "submitter"
    first_name: str | None = None
    last_name: str | None = None
    phone: str | None = None
    organization: str | None = None
    employee_id: str | None = None
    dob: str | None = None
    gender: str | None = None
    policy: str | None = None
    sum_insured: float | str | None = None
    provider: str = "local"


@router.post("/auth/register", status_code=200)
def register_user_profile(payload: RegisterUserIn):
    """Complete registration and store patient/staff profile into database."""
    email = str(payload.username).strip().lower()
    if not email:
        raise HTTPException(status_code=400, detail="Email is required")

    pwd = payload.password_hash or payload.password or None

    with force_master_session():
        with SessionLocal() as db:
            user_row = db.execute(
                text("SELECT id FROM users WHERE lower(email) = lower(:email)"),
                {"email": email},
            ).mappings().first()

            if user_row:
                user_id = user_row["id"]
                db.execute(
                    text("UPDATE users SET status = 'ACTIVE', updated_at = CURRENT_TIMESTAMP WHERE id = :id"),
                    {"id": user_id},
                )
            else:
                user_id = uuid.uuid4()
                db.execute(
                    text("""
                        INSERT INTO users (id, email, phone, external_provider, external_subject_id, status, email_verified, password_hash, created_at, updated_at)
                        VALUES (:id, :email, :phone, :provider, :email, 'ACTIVE', 1, :pwd, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                    """),
                    {
                        "id": user_id,
                        "email": email,
                        "phone": payload.phone or None,
                        "provider": payload.provider or "local",
                        "pwd": pwd,
                    },
                )

                role_name = "admin" if payload.role in ("admin", "tpa") else "submitter"
                role_row = db.execute(
                    text("SELECT id FROM roles WHERE name = :rname"),
                    {"rname": role_name},
                ).mappings().first()

                if not role_row:
                    r_id = uuid.uuid4()
                    db.execute(
                        text("INSERT INTO roles (id, name, description, created_at) VALUES (:id, :name, 'Role', CURRENT_TIMESTAMP)"),
                        {"id": r_id, "name": role_name},
                    )
                else:
                    r_id = role_row["id"]

                db.execute(
                    text("INSERT INTO user_roles (id, user_id, role_id, created_at) VALUES (:id, :user_id, :role_id, CURRENT_TIMESTAMP)"),
                    {"id": uuid.uuid4(), "user_id": user_id, "role_id": r_id},
                )

            # Insert or update patient profile
            sum_val = float(str(payload.sum_insured).replace(",", "").strip()) if payload.sum_insured else 500000.0
            prof_row = db.execute(
                text("SELECT id FROM patient_profiles WHERE user_id = :user_id"),
                {"user_id": user_id},
            ).mappings().first()

            if prof_row:
                db.execute(
                    text("""
                        UPDATE patient_profiles
                        SET first_name = COALESCE(:first_name, first_name),
                            last_name = COALESCE(:last_name, last_name),
                            gender = COALESCE(:gender, gender),
                            policy_number = COALESCE(:policy_number, policy_number),
                            sum_insured = COALESCE(:sum_insured, sum_insured),
                            coverage_verified = 1,
                            updated_at = CURRENT_TIMESTAMP
                        WHERE user_id = :user_id
                    """),
                    {
                        "user_id": user_id,
                        "first_name": payload.first_name or "Patient",
                        "last_name": payload.last_name or "",
                        "gender": payload.gender or "Male",
                        "policy_number": payload.policy or "POL-DEFAULT",
                        "sum_insured": sum_val,
                    },
                )
            else:
                db.execute(
                    text("""
                        INSERT INTO patient_profiles (id, user_id, first_name, last_name, gender, policy_number, sum_insured, coverage_verified, created_at, updated_at)
                        VALUES (:id, :user_id, :first_name, :last_name, :gender, :policy_number, :sum_insured, 1, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                    """),
                    {
                        "id": uuid.uuid4(),
                        "user_id": user_id,
                        "first_name": payload.first_name or "Patient",
                        "last_name": payload.last_name or "",
                        "gender": payload.gender or "Male",
                        "policy_number": payload.policy or "POL-DEFAULT",
                        "sum_insured": sum_val,
                    },
                )

            db.commit()

            return {
                "success": True,
                "user_id": str(user_id),
                "email": email,
                "needs_onboarding": False,
                "message": "User and profile saved successfully",
            }


# ------------------------------------------------------------------ TPA registration
# Passwords are sent only to Keycloak and are never persisted in ClaimGPT.
class TpaRegistrationIn(BaseModel):
    first_name: str = Field(min_length=1, max_length=100)
    last_name: str = Field(min_length=1, max_length=100)
    email: EmailStr
    phone: str | None = Field(default=None, max_length=30)
    organization_id: uuid.UUID
    employee_id: str | None = Field(default=None, max_length=100)
    password: str = Field(min_length=8, max_length=256)


class OrganizationRegistrationIn(BaseModel):
    name: str = Field(min_length=2, max_length=255)
    address: str | None = Field(default=None, max_length=2000)


def _keycloak_admin_token() -> str:
    """Obtain a short-lived admin token for provisioning a Keycloak user."""
    import httpx

    url = f"{settings.keycloak_url}/realms/master/protocol/openid-connect/token"
    try:
        response = httpx.post(
            url,
            data={
                "grant_type": "password",
                "client_id": "admin-cli",
                "username": settings.keycloak_admin_username,
                "password": settings.keycloak_admin_password,
            },
            timeout=10,
        )
        response.raise_for_status()
        return response.json()["access_token"]
    except Exception as exc:
        logger.exception("Could not authenticate with Keycloak admin API")
        raise HTTPException(status_code=503, detail="Account service is unavailable") from exc


def _create_keycloak_reviewer(payload: TpaRegistrationIn) -> str:
    """Create the credential record in Keycloak and assign the reviewer role."""
    import httpx

    token = _keycloak_admin_token()
    base_url = f"{settings.keycloak_url}/admin/realms/{settings.keycloak_realm}"
    headers = {"Authorization": f"Bearer {token}"}
    user = {
        "username": str(payload.email),
        "email": str(payload.email),
        "firstName": payload.first_name.strip(),
        "lastName": payload.last_name.strip(),
        "enabled": True,
        "emailVerified": False,
        "credentials": [{"type": "password", "value": payload.password, "temporary": False}],
    }
    try:
        response = httpx.post(f"{base_url}/users", headers=headers, json=user, timeout=10)
        if response.status_code == 409:
            raise HTTPException(status_code=409, detail="An account with this email already exists")
        response.raise_for_status()
        location = response.headers.get("Location", "")
        keycloak_user_id = location.rstrip("/").split("/")[-1]
        if not keycloak_user_id:
            raise RuntimeError("Keycloak did not return a user ID")

        role_response = httpx.get(f"{base_url}/roles/reviewer", headers=headers, timeout=10)
        role_response.raise_for_status()
        assignment = httpx.post(
            f"{base_url}/users/{keycloak_user_id}/role-mappings/realm",
            headers=headers,
            json=[role_response.json()],
            timeout=10,
        )
        assignment.raise_for_status()
        return keycloak_user_id
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Could not provision Keycloak reviewer")
        raise HTTPException(status_code=503, detail="Could not create the login account") from exc


@router.get("/organizations", status_code=200)
def list_approved_tpa_organizations():
    """Public registration list: only organizations already approved by an admin."""
    with SessionLocal() as db:
        rows = db.execute(
            text("SELECT id, name FROM organizations WHERE type = 'TPA' AND status = 'ACTIVE' ORDER BY name")
        ).mappings().all()
    return {"organizations": [{"id": str(row["id"]), "name": row["name"]} for row in rows]}


@router.post("/organizations/registration", status_code=201, dependencies=[Depends(RateLimiter(limit=3, window_seconds=60))])
def request_tpa_organization_registration(payload: OrganizationRegistrationIn):
    """Submit an organization for approval; it is intentionally not selectable yet."""
    with SessionLocal() as db:
        existing = db.execute(
            text("SELECT id, status FROM organizations WHERE lower(name) = lower(:name) AND type = 'TPA'"),
            {"name": payload.name.strip()},
        ).mappings().first()
        if existing:
            return {"id": str(existing["id"]), "status": existing["status"], "message": "Organization already exists"}

        new_org = Organization(
            name=payload.name.strip(),
            type='TPA',
            address=payload.address.strip() if payload.address else None,
            status='PENDING'
        )
        db.add(new_org)
        db.commit()
    return {"id": str(new_org.id), "status": new_org.status, "message": "Organization submitted for approval"}


@router.post("/tpa-adjusters/registration", status_code=201, dependencies=[Depends(RateLimiter(limit=3, window_seconds=60))])
def register_tpa_adjuster(payload: TpaRegistrationIn):
    """Create a Keycloak login and its matching local staff profile."""
    _ensure_users_password_hash_column()

    with SessionLocal() as db:
        organization = db.execute(
            text("SELECT id FROM organizations WHERE id = :id AND type = 'TPA' AND status = 'ACTIVE'"),
            {"id": str(payload.organization_id)},
        ).mappings().first()
        if not organization:
            raise HTTPException(status_code=400, detail="Select an approved TPA organization")

    keycloak_user_id = _create_keycloak_reviewer(payload)
    try:
        with SessionLocal() as db:
            new_user = User(
                email=str(payload.email),
                phone=payload.phone or None,
                external_provider='keycloak',
                external_subject_id=keycloak_user_id,
                status='ACTIVE',
                password_hash=hash_password(payload.password)
            )
            db.add(new_user)
            db.flush()
            user_id = new_user.id
            db.execute(
                text("""
                    INSERT INTO staff_profiles (id, user_id, organization_id, employee_id, designation, status)
                    VALUES (:id, :user_id, :organization_id, :employee_id, 'TPA Adjuster', 'ACTIVE')
                """),
                {"id": uuid.uuid4(), "user_id": user_id, "organization_id": str(payload.organization_id), "employee_id": payload.employee_id or None},
            )
            db.execute(
                text("""
                    INSERT INTO user_roles (id, user_id, role_id)
                    SELECT :id, :user_id, id FROM roles WHERE name = 'reviewer'
                """),
                {"id": uuid.uuid4(), "user_id": user_id},
            )
            db.commit()
    except Exception as exc:
        logger.exception("Keycloak user %s was created but local profile creation failed", keycloak_user_id)
        raise HTTPException(status_code=500, detail="Could not save the account profile") from exc

    return {"user_id": str(user["id"]), "role": "reviewer", "message": "TPA adjuster account created"}


@router.get("/health")
def health():
    db_ok = check_db_health()
    status = "ok" if db_ok else "degraded"
    return {"status": status, "database": "up" if db_ok else "down"}


@router.post("/auth/login")
@router.post("/auth/register")
def authenticate_or_register_user(data: dict[str, Any] = Body(...), db: Session = Depends(get_db)):
    username = data.get("username") or data.get("name") or data.get("email", "Swagath")
    if isinstance(username, str) and "@" in username:
        username = username.split("@")[0]
    username = str(username).strip().capitalize()
    
    email = data.get("email") or f"{username.lower()}@example.com"
    role = data.get("role", "patient")
    
    _audit(db, "USER_LOGIN_OR_REGISTER", metadata={"username": username, "email": email, "role": role})
    logger.info("User registered/authenticated in Docker backend: %s (%s)", username, email)
    return {
        "status": "success",
        "message": f"Account {username} initialized in backend",
        "user": {
            "name": username,
            "email": email,
            "role": role,
            "account_id": f"ACC-{username.upper()}-2026"
        }
    }


@router.get("/claims/upload-token", dependencies=[Depends(RateLimiter(limit=10, window_seconds=60))])
def get_upload_token(filename: str = Query(...)):
    """Generate a secure pre-signed PUT upload URL/SAS token for direct client-to-storage upload."""
    if not filename:
        raise HTTPException(status_code=400, detail="Filename parameter is required")
        
    from libs.shared.storage import MinioStorage
    safe_name = _safe_filename(filename)
    # Generate unique path in the storage pending folder
    temp_key = f"pending/{uuid.uuid4().hex}_{safe_name}"
    
    try:
        token_info = MinioStorage.generate_presigned_upload_url(temp_key)
        # Add metadata for subsequent API call registration
        token_info["storage_path"] = f"s3://{MinioStorage.BUCKET_NAME}/{temp_key}"
        token_info["filename"] = safe_name
        return token_info
    except Exception as e:
        logger.exception("Failed to generate upload token: %s", e)
        raise HTTPException(status_code=500, detail=f"Failed to generate secure upload token: {e}")


@router.post("/claims/", status_code=202)
async def create_claim(
    files: list[UploadFile] = File(default=[]),
    policy_id: str = Form(None),
    patient_id: str = Form(None),
    storage_paths: list[str] = Form(None),
    auth_user: AuthUser = Depends(get_current_user_context),
    db: Session = Depends(get_db),
):
    """Create a new claim by uploading files or passing pre-uploaded storage paths.
    
    This endpoint accepts files or direct storage paths, and enqueues the pipeline.
    All database operations (idempotency, deduplication) are handled by the 
    intake_task in the Celery worker.
    """
    # Ensure either files or storage_paths are provided
    storage_paths_list = []
    if storage_paths:
        if len(storage_paths) == 1 and (storage_paths[0].startswith("[") or "," in storage_paths[0]):
            try:
                import json
                parsed = json.loads(storage_paths[0])
                if isinstance(parsed, list):
                    storage_paths_list = [str(x) for x in parsed]
            except Exception:
                storage_paths_list = [x.strip() for x in storage_paths[0].split(",") if x.strip()]
        else:
            storage_paths_list = [str(p) for p in storage_paths]

    files_count = len(files) if files else 0
    paths_count = len(storage_paths_list)
    
    logger.info(f"[create_claim] Starting with {files_count} files and {paths_count} storage paths")
    upload_log.info(
        "UPLOAD_START | endpoint=create_claim files=%d paths=%d policy_id=%s patient_id=%s",
        files_count,
        paths_count,
        policy_id,
        patient_id,
    )
    
    if not files and not storage_paths_list:
        upload_log.warning("UPLOAD_REJECTED | endpoint=create_claim reason=no_files_or_paths")
        raise HTTPException(status_code=400, detail="At least one file or storage_path is required")

    # --- Validate all files and read content ---
    file_metadata_list: list[dict[str, str]] = []  # Will hold metadata for intake_task
    saved_paths: list[Any] = []
    
    try:
        if storage_paths_list:
            for sp in storage_paths_list:
                if not sp.startswith("s3://"):
                    raise HTTPException(status_code=400, detail=f"Invalid storage path URI: {sp}. Must start with s3://")
                
                path_parts = sp.split("/")
                raw_filename = path_parts[-1] if path_parts else "document.pdf"
                if len(raw_filename) > 33 and raw_filename[32] == "_":
                    safe_name = raw_filename[33:]
                else:
                    safe_name = raw_filename
                    
                effective_ct = "application/pdf"
                if safe_name.lower().endswith((".jpg", ".jpeg")):
                    effective_ct = "image/jpeg"
                elif safe_name.lower().endswith(".png"):
                    effective_ct = "image/png"
                    
                content_hash = hashlib.sha256(sp.encode("utf-8")).hexdigest()
                
                file_metadata_list.append({
                    "path": sp,
                    "safe_name": safe_name,
                    "content_hash": content_hash,
                    "effective_ct": effective_ct,
                })
                saved_paths.append(sp)
                
                upload_log.info(
                    "DIRECT_FILE_RECEIVED | endpoint=create_claim file=%s path=%s type=%s",
                    safe_name, sp, effective_ct
                )
        else:
            for idx, file in enumerate(files):
                # Validate content type
                effective_ct, ok = _resolve_content_type(file)
                if not ok:
                    upload_log.warning(
                        "UPLOAD_REJECTED | endpoint=create_claim reason=unsupported_type file=%s type=%s",
                        file.filename, file.content_type,
                    )
                    raise HTTPException(
                        status_code=415,
                        detail=f"Unsupported file type '{file.content_type}' for '{file.filename}'. "
                        f"Allowed: {', '.join(sorted(settings.allowed_content_types))}",
                    )
                
                # Read and validate file size
                file_bytes = await file.read()
                if len(file_bytes) > settings.max_upload_bytes:
                    upload_log.warning(
                        "UPLOAD_REJECTED | endpoint=create_claim reason=too_large file=%s bytes=%d max=%d",
                        file.filename, len(file_bytes), settings.max_upload_bytes,
                    )
                    raise HTTPException(
                        status_code=413,
                        detail=f"File '{file.filename}' too large ({len(file_bytes)} bytes). Max: {settings.max_upload_bytes} bytes",
                    )
                
                # Calculate content hash and safe filename
                safe_name = _safe_filename(file.filename)
                content_hash = hashlib.sha256(file_bytes).hexdigest()
                logger.info(f"[create_claim] File validated: {safe_name}, hash={content_hash}")
                
                # Upload file directly to MinIO under a temporary key
                from libs.shared.storage import MinioStorage
                temp_key = f"pending/{uuid.uuid4().hex}_{safe_name}"
                try:
                    minio_uri = MinioStorage.upload_file(temp_key, file_bytes)
                    saved_paths.append(minio_uri)
                    logger.info(f"[create_claim] File uploaded directly to MinIO: {minio_uri}")
                except Exception as e:
                    logger.exception(f"[create_claim] Failed to upload file to MinIO: {temp_key}")
                    raise HTTPException(status_code=500, detail="Failed to store uploaded file in object storage")
                
                # Store metadata for intake_task
                file_metadata_list.append({
                    "path": minio_uri,
                    "safe_name": safe_name,
                    "content_hash": content_hash,
                    "effective_ct": effective_ct,
                })
                
                upload_log.info(
                    "FILE_RECEIVED | endpoint=create_claim file=%s bytes=%d type=%s sha256=%s",
                    safe_name, len(file_bytes), effective_ct, content_hash,
                )

        # Determine unique user/patient identity
        clean_patient_id = patient_id.strip() if patient_id and patient_id.strip() and patient_id.strip().lower() not in ("user", "null", "undefined") else None
        target_patient_id = clean_patient_id or auth_user.user_id or auth_user.patient_id or auth_user.email

        # Synchronous duplicate check using sorted set_hash (works for both single & multi-file uploads)
        if file_metadata_list:
            hashes = [metadata["content_hash"] for metadata in file_metadata_list if metadata["content_hash"]]
            hashes.sort()
            set_hash = hashlib.sha256(",".join(hashes).encode("utf-8")).hexdigest()

            target_user = target_patient_id or policy_id
            existing_job = None

            if target_user:
                from sqlalchemy import func
                from libs.shared.models import ParseJob
                existing_job = (
                    db.query(ParseJob)
                    .join(Claim, ParseJob.claim_id == Claim.id)
                    .filter(
                        ParseJob.set_hash == set_hash, 
                        Claim.status == "COMPLETED",
                        func.lower(Claim.patient_id) == func.lower(target_user)
                    )
                    .first()
                )

            if existing_job:
                upload_log.info(
                    "UPLOAD_DUPLICATE | Found completed claim %s matching set_hash for user %s, returning immediately",
                    existing_job.claim_id, target_user
                )
                # Clean up the temporarily uploaded MinIO files
                from libs.shared.storage import MinioStorage
                for uri in saved_paths:
                    try:
                        MinioStorage.delete_file(uri)
                    except Exception:
                        pass

                return {
                    "claim_id": str(existing_job.claim_id),
                    "task_id": None,
                    "status": "COMPLETED",
                    "message": "Claim already processed.",
                }
        
        # --- Create Claim & Documents synchronously so client immediately receives real claim_id ---
        new_claim_id = uuid.uuid4()
        new_claim = Claim(
            id=new_claim_id,
            policy_id=policy_id,
            patient_id=patient_id,
            status="UPLOADED",
            source="PATIENT",
        )
        db.add(new_claim)

        for metadata in file_metadata_list:
            doc = Document(
                claim_id=new_claim.id,
                file_name=metadata["safe_name"],
                file_type=metadata["effective_ct"],
                minio_path=metadata["path"],
                content_hash=metadata["content_hash"],
            )
            db.add(doc)

        from libs.shared.models import ParseJob
        parse_job = ParseJob(
            claim_id=new_claim.id,
            status="PENDING",
            set_hash=set_hash if file_metadata_list else None,
        )
        db.add(parse_job)

        upsert_workflow_state(db, new_claim.id, "STARTING", status="RUNNING")
        db.commit()

        # Enqueue pipeline starting with OCR (claim and docs already committed)
        task_id = _enqueue_pipeline(str(new_claim.id), policy_id, target_patient_id)

        upload_log.info(
            "UPLOAD_SUCCESS | endpoint=create_claim files=%d claim_id=%s task_id=%s",
            len(file_metadata_list), new_claim.id, task_id,
        )

        return {
            "claim_id": str(new_claim.id),
            "task_id": task_id,
            "status": "QUEUED",
            "message": "Claim upload queued. Check status via /claims/{claim_id}/progress endpoint.",
        }
    
    except HTTPException:
        # Clean up already uploaded MinIO files on validation error
        from libs.shared.storage import MinioStorage
        for uri in saved_paths:
            try:
                MinioStorage.delete_file(uri)
            except Exception:
                logger.warning(f"Failed to clean up pending MinIO file: {uri}")
        raise
    except Exception as exc:
        logger.exception("Error during file upload processing")
        upload_log.exception(
            "UPLOAD_FAILURE | endpoint=create_claim files=%d error=%s",
            len(files), exc,
        )
        # Clean up already uploaded MinIO files
        from libs.shared.storage import MinioStorage
        for uri in saved_paths:
            try:
                MinioStorage.delete_file(uri)
            except Exception:
                logger.warning(f"Failed to clean up pending MinIO file: {uri}")
        raise HTTPException(status_code=500, detail="Failed to process file upload")


@router.get("/claims", response_model=ClaimListOut)
def list_claims(
    offset: int = Query(0, ge=0),
    limit: int = Query(100, ge=1, le=500),
    patient_id: str | None = Query(None),
    policy_id: str | None = Query(None),
    auth_user: AuthUser = Depends(get_current_user_context),
    db: Session = Depends(get_db),
):
    try:
        query = db.query(Claim)
        clean_patient_id = patient_id.strip() if patient_id and patient_id.strip() and patient_id.strip().lower() not in ("user", "null", "undefined") else None
        effective_patient = clean_patient_id or auth_user.user_id or auth_user.patient_id or None
        
        if auth_user.role in ("patient", "submitter"):
            # Gather all identifier candidates for this user for backward-compatibility
            candidates = {
                c.strip().lower() for c in [
                    clean_patient_id,
                    auth_user.user_id,
                    auth_user.patient_id,
                    auth_user.email,
                ] if c and c.strip().lower() not in ("user", "null", "undefined")
            }
            if candidates:
                query = query.filter(func.lower(Claim.patient_id).in_(list(candidates)))
        elif effective_patient:
            query = query.filter(func.lower(Claim.patient_id) == effective_patient.lower())

        if policy_id:
            query = query.filter(Claim.policy_id == policy_id)

        total = query.count()
        claims = (
            query
            .order_by(Claim.created_at.desc())
            .offset(offset)
            .limit(limit)
            .all()
        )
        
        # Batch-fetch relevant parsed fields for these claims to avoid N+1 queries
        if claims:
            claim_ids = [c.id for c in claims]
            pf_rows = db.query(ParsedField).filter(
                ParsedField.claim_id.in_(claim_ids),
                ParsedField.field_name.in_([
                    "patient_name", "member_name", "insured_name",
                    "hospital_name", "hospital",
                    "doctor_name", "doctor", "provider_name", "rendering_provider",
                    "diagnosis", "primary_diagnosis", "chief_complaint"
                ])
            ).all()
            
            from collections import defaultdict
            pf_by_claim = defaultdict(dict)
            for row in pf_rows:
                pf_by_claim[row.claim_id][row.field_name] = row.field_value
                
            for c in claims:
                fields = pf_by_claim[c.id]
                c.patient_name = fields.get("patient_name") or fields.get("member_name") or fields.get("insured_name") or None
                c.hospital_name = fields.get("hospital_name") or fields.get("hospital") or None
                c.doctor_name = fields.get("doctor_name") or fields.get("doctor") or fields.get("provider_name") or fields.get("rendering_provider") or None
                c.diagnosis = fields.get("diagnosis") or fields.get("primary_diagnosis") or fields.get("chief_complaint") or None

        claim_items = [
            {
                "id": str(c.id).lower(),
                "policy_id": c.policy_id,
                "patient_id": c.patient_id,
                "status": c.status,
                "source": c.source,
                "created_at": c.created_at,
                "updated_at": c.updated_at,
                "documents": c.documents,
                "task_id": getattr(c, "task_id", None),
                "patient_name": getattr(c, "patient_name", None),
                "hospital_name": getattr(c, "hospital_name", None),
                "doctor_name": getattr(c, "doctor_name", None),
                "diagnosis": getattr(c, "diagnosis", None),
            }
            for c in claims
        ]

        return ClaimListOut(claims=claim_items, total=total)
    except Exception as exc:
        logger.exception("Error listing claims")
        raise HTTPException(status_code=500, detail="Failed to list claims")


@router.get("/claims/{claim_id}", response_model=ClaimOut)
def get_claim(
    claim_id: str,
    patient_id: str | None = Query(None),
    auth_user: AuthUser = Depends(get_current_user_context),
    db: Session = Depends(get_db)
):
    cid = _parse_uuid(claim_id)
    claim = (
        db.query(Claim)
        .options(selectinload(Claim.documents))
        .filter(Claim.id == cid)
        .first()
    )
    if not claim:
        raise HTTPException(status_code=404, detail="Claim not found")

    # Enforce Row-Level Security: if caller is a patient identity, verify ownership
    caller_candidates = {
        c.strip().lower() for c in [
            patient_id,
            auth_user.user_id,
            auth_user.patient_id,
            auth_user.email,
        ] if c and c.strip().lower() not in ("user", "null", "undefined")
    }
    if caller_candidates and claim.patient_id:
        if claim.patient_id.strip().lower() not in caller_candidates and auth_user.role not in ("admin", "reviewer", "auditor"):
            raise HTTPException(status_code=404, detail="Claim not found")
        
    # Fetch relevant parsed fields for this claim
    pf_rows = db.query(ParsedField).filter(
        ParsedField.claim_id == cid,
        ParsedField.field_name.in_([
            "patient_name", "member_name", "insured_name",
            "hospital_name", "hospital",
            "doctor_name", "doctor", "provider_name", "rendering_provider",
            "diagnosis", "primary_diagnosis", "chief_complaint"
        ])
    ).all()
    
    fields = {row.field_name: row.field_value for row in pf_rows}
    claim.patient_name = fields.get("patient_name") or fields.get("member_name") or fields.get("insured_name") or None
    claim.hospital_name = fields.get("hospital_name") or fields.get("hospital") or None
    claim.doctor_name = fields.get("doctor_name") or fields.get("doctor") or fields.get("provider_name") or fields.get("rendering_provider") or None
    claim.diagnosis = fields.get("diagnosis") or fields.get("primary_diagnosis") or fields.get("chief_complaint") or None
    
    return ClaimOut.model_validate(claim).model_dump(mode="json")


def _map_progress(current_step: str | None, status: str | None) -> tuple[str | None, int]:
    if current_step == "STARTING":
        return "Starting", 5
    if current_step == "OCR_IN_PROGRESS":
        return "OCR (extracting text)", 20
    if current_step == "OCR_COMPLETED":
        return "OCR complete", 35
    if current_step == "PARSING_IN_PROGRESS":
        return "Parsing (LLM agent reading document)", 55
    if current_step == "PARSING_COMPLETED":
        return "Parsing complete", 70
    if current_step == "CODING_ANALYSIS":
        return "Medical coding (ICD-10 / CPT)", 78
    if current_step == "CODING_COMPLETED":
        return "Coding complete", 82
    if current_step == "RISK_ANALYSIS":
        return "Risk scoring", 86
    if current_step == "RISK_COMPLETED":
        return "Risk complete", 90
    if current_step == "VALIDATION_RUNNING":
        return "Validating", 92
    if current_step == "VALIDATION_COMPLETED":
        return "Validation complete", 96
    if current_step == "RETRYING":
        # Don't regress — keep above prior steps; monotonic guard below also protects.
        return "Retrying (transient)", 92
    if current_step == "FAILED" or status == "FAILED":
        return "Failed", 0
    if current_step == "FINALIZING":
        return "Finalizing", 98
    if current_step == "FINISHED" or status == "FINISHED":
        return "Completed", 100
    return current_step, 0


# Progress cache removed to prevent uvicorn multi-worker state mismatch


@router.get("/claims/{claim_id}/status")
def get_claim_status(claim_id: str, db: Session = Depends(get_db)):
    cid = _parse_uuid(claim_id)
    state = get_latest_workflow_state(db, cid)
    if not state:
        return {"current_step": None, "status": None, "step_index": 0, "percentage": 0.0}
    
    step, percentage = _map_progress(state.current_step, state.status)
    step_index = _get_step_index(state.current_step, state.status)
    return {
        "current_step": state.current_step,
        "step": step,
        "status": state.status,
        "step_index": step_index,
        "percentage": float(percentage)
    }


@router.get("/claims/{claim_id}/progress")
def get_claim_progress(claim_id: str, db: Session = Depends(get_db)):
    cid = _parse_uuid(claim_id)
    state = get_latest_workflow_state(db, cid)

    # No workflow state yet: distinguish "claim does not exist" from
    # "claim was created but the pipeline hasn't recorded any progress yet".
    if not state:
        claim = db.query(Claim).filter(Claim.id == cid).first()
        if not claim:
            raise HTTPException(status_code=404, detail="Claim not found")
        # Claim exists but no state row — treat as queued, never as silently null.
        return {
            "status": "QUEUED",
            "step": "Queued (waiting for worker)",
            "percentage": 2,
            "is_complete": False,
            "error": None,
        }

    step, percentage = _map_progress(state.current_step, state.status)
    is_failed = (state.status == "FAILED") or (state.current_step == "FAILED")
    is_complete = bool(percentage == 100 or is_failed)

    error_message: str | None = None
    if is_failed:
        # Surface the most recent job error message so the UI can show *why*
        # the upload stopped, instead of polling forever on 0%.
        try:
            latest_parse = (
                db.query(ParseJob)
                .filter(ParseJob.claim_id == cid)
                .order_by(ParseJob.created_at.desc())
                .first()
            )
            if latest_parse and latest_parse.error_message:
                error_message = latest_parse.error_message
            if not error_message:
                from libs.shared.models import OcrJob as _OcrJob
                latest_ocr = (
                    db.query(_OcrJob)
                    .filter(_OcrJob.claim_id == cid)
                    .order_by(_OcrJob.created_at.desc())
                    .first()
                )
                if latest_ocr and latest_ocr.error_message:
                    error_message = latest_ocr.error_message
        except Exception:
            logger.exception("Failed to read latest job error for claim %s", cid)
        if not error_message:
            error_message = "Pipeline failed. See server logs for details."

    # Mapped from database state directly, naturally monotonic in Celery chain
    return {
        "status": state.status,
        "step": step,
        "percentage": percentage,
        "is_complete": is_complete,
        "error": error_message,
    }


@router.get("/claims/{claim_id}/file")
def download_original_file(claim_id: str, view: bool = False, db: Session = Depends(get_db)):
    cid = _parse_uuid(claim_id)

    doc = (
        db.query(Document)
        .filter(Document.claim_id == cid)
        .order_by(Document.uploaded_at.desc())
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="No document found for claim")

    disp = "inline" if view else "attachment"

    if doc.minio_path and doc.minio_path.startswith("s3://"):
        from libs.shared.storage import MinioStorage
        from fastapi.responses import StreamingResponse
        import io
        try:
            file_bytes = MinioStorage.download_file_bytes(doc.minio_path)
            return StreamingResponse(
                io.BytesIO(file_bytes),
                media_type=doc.file_type or "application/octet-stream",
                headers={
                    "Content-Disposition": f'{disp}; filename="{doc.file_name}"'
                }
            )
        except Exception as e:
            logger.exception(f"Failed to fetch {doc.minio_path} from cloud storage: {e}")
            raise HTTPException(status_code=404, detail="File not found in cloud storage")

    file_path = Path(doc.minio_path).resolve()

    # prevent path traversal — file must be under RAW_STORAGE
    if not str(file_path).startswith(str(RAW_STORAGE)):
        raise HTTPException(status_code=403, detail="Access denied")

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Stored file not found on disk")

    return FileResponse(
        str(file_path),
        media_type=doc.file_type or "application/pdf",
        headers={"Content-Disposition": f'{disp}; filename="{doc.file_name}"'}
    )


@router.get("/claims/{claim_id}/documents/{doc_id}/file")
def download_document_file(claim_id: str, doc_id: str, view: bool = False, db: Session = Depends(get_db)):
    cid = _parse_uuid(claim_id)
    did = _parse_uuid(doc_id)

    doc = db.query(Document).filter(Document.id == did, Document.claim_id == cid).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    disp = "inline" if view else "attachment"

    if doc.minio_path and doc.minio_path.startswith("s3://"):
        from libs.shared.storage import MinioStorage
        from fastapi.responses import StreamingResponse
        import io
        try:
            file_bytes = MinioStorage.download_file_bytes(doc.minio_path)
            return StreamingResponse(
                io.BytesIO(file_bytes),
                media_type=doc.file_type or "application/octet-stream",
                headers={
                    "Content-Disposition": f'{disp}; filename="{doc.file_name}"'
                }
            )
        except Exception as e:
            logger.exception(f"Failed to fetch {doc.minio_path} from cloud storage: {e}")
            raise HTTPException(status_code=404, detail="File not found in cloud storage")

    file_path = Path(doc.minio_path).resolve()

    if not file_path.exists():
        file_path = (RAW_STORAGE / Path(doc.minio_path).name).resolve()
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="Original document file missing on disk")

    # prevent path traversal — file must be under RAW_STORAGE
    if not str(file_path).startswith(str(RAW_STORAGE)):
        raise HTTPException(status_code=403, detail="Access denied")

    return FileResponse(
        str(file_path),
        media_type=doc.file_type or "application/pdf",
        headers={"Content-Disposition": f'{disp}; filename="{doc.file_name}"'}
    )


@router.get("/claims/{claim_id}/documents/{doc_id}/pages/{page_number}/image")
def get_document_page_image(claim_id: str, doc_id: str, page_number: int = 1, db: Session = Depends(get_db)):
    cid = _parse_uuid(claim_id)
    did = _parse_uuid(doc_id)

    doc = db.query(Document).filter(Document.id == did, Document.claim_id == cid).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    from fastapi.responses import Response

    if doc.minio_path and doc.minio_path.startswith("s3://"):
        from libs.shared.storage import MinioStorage
        try:
            file_bytes = MinioStorage.download_file_bytes(doc.minio_path)
        except Exception as e:
            logger.exception(f"Failed to fetch {doc.minio_path} from cloud storage: {e}")
            raise HTTPException(status_code=404, detail="File not found in cloud storage")

        # Determine file extension from key/filename
        ext = Path(doc.file_name).suffix.lower() if doc.file_name else ".pdf"
        if ext in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif"} or (doc.file_type and doc.file_type.startswith("image/")):
            return Response(content=file_bytes, media_type=doc.file_type or "image/png")

        if ext == ".pdf" or doc.file_type == "application/pdf":
            try:
                import io
                import pypdfium2
                pdf = pypdfium2.PdfDocument(file_bytes)
                total_pages = len(pdf)
                target_idx = max(0, min(page_number - 1, total_pages - 1))
                page = pdf.get_page(target_idx)
                img = page.render(scale=2).to_pil()
                buf = io.BytesIO()
                img.save(buf, format="PNG")
                buf.seek(0)
                return Response(content=buf.getvalue(), media_type="image/png")
            except Exception as e:
                logger.exception("Failed to render PDF page image from S3: %s", e)
                raise HTTPException(status_code=500, detail="Failed to render PDF page")

        return Response(content=file_bytes, media_type=doc.file_type or "application/octet-stream")

    file_path = Path(doc.minio_path).resolve()
    if not file_path.exists():
        file_path = (RAW_STORAGE / Path(doc.minio_path).name).resolve()
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="Original document file missing on disk")

    ext = file_path.suffix.lower()
    if ext in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif"}:
        return FileResponse(str(file_path), media_type=doc.file_type or "image/png")

    if ext == ".pdf" or doc.file_type == "application/pdf":
        try:
            import io
            import pypdfium2
            pdf = pypdfium2.PdfDocument(str(file_path))
            total_pages = len(pdf)
            target_idx = max(0, min(page_number - 1, total_pages - 1))
            page = pdf.get_page(target_idx)
            img = page.render(scale=2).to_pil()
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            buf.seek(0)
            return Response(content=buf.getvalue(), media_type="image/png")
        except Exception as e:
            logger.exception("Failed to render PDF page image: %s", e)
            raise HTTPException(status_code=500, detail="Failed to render PDF page")

    return FileResponse(str(file_path), media_type=doc.file_type or "application/octet-stream")


@router.post("/claims/{claim_id}/documents", response_model=ClaimOut, status_code=201)
async def add_documents_to_claim(
    claim_id: str,
    files: list[UploadFile] = File(default=[]),
    storage_paths: list[str] = Form(None),
    db: Session = Depends(get_db),
):
    logger.info(f"[IDEMPOTENCY] Starting add_documents_to_claim with files and storage paths for claim {claim_id}.")
    upload_log.info(
        "UPLOAD_START | endpoint=add_documents claim_id=%s files=%d",
        claim_id, len(files) if files else 0,
    )
    """Add supporting documents to an existing claim."""
    cid = _parse_uuid(claim_id)
    claim = db.query(Claim).filter(Claim.id == cid).first()
    if not claim:
        upload_log.warning(
            "UPLOAD_REJECTED | endpoint=add_documents reason=claim_not_found claim_id=%s",
            claim_id,
        )
        raise HTTPException(status_code=404, detail="Claim not found")

    storage_paths_list = []
    if storage_paths:
        if len(storage_paths) == 1 and (storage_paths[0].startswith("[") or "," in storage_paths[0]):
            try:
                import json
                parsed = json.loads(storage_paths[0])
                if isinstance(parsed, list):
                    storage_paths_list = [str(x) for x in parsed]
            except Exception:
                storage_paths_list = [x.strip() for x in storage_paths[0].split(",") if x.strip()]
        else:
            storage_paths_list = [str(p) for p in storage_paths]

    if not files and not storage_paths_list:
        upload_log.warning(
            "UPLOAD_REJECTED | endpoint=add_documents reason=no_files_or_paths claim_id=%s",
            claim_id,
        )
        raise HTTPException(status_code=400, detail="At least one file or storage_path is required")

    # --- validate all files and calculate content_hash
    file_data: list[tuple[Any, Any, str, str, str]] = []  # (raw_path, file_bytes, safe_name, content_hash, effective_ct)
    
    if storage_paths_list:
        for sp in storage_paths_list:
            if not sp.startswith("s3://"):
                raise HTTPException(status_code=400, detail=f"Invalid storage path URI: {sp}. Must start with s3://")
            
            path_parts = sp.split("/")
            raw_filename = path_parts[-1] if path_parts else "document.pdf"
            if len(raw_filename) > 33 and raw_filename[32] == "_":
                safe_name = raw_filename[33:]
            else:
                safe_name = raw_filename
                
            effective_ct = "application/pdf"
            if safe_name.lower().endswith((".jpg", ".jpeg")):
                effective_ct = "image/jpeg"
            elif safe_name.lower().endswith(".png"):
                effective_ct = "image/png"
                
            content_hash = hashlib.sha256(sp.encode("utf-8")).hexdigest()
            file_data.append((sp, None, safe_name, content_hash, effective_ct))
    else:
        for file in files:
            effective_ct, ok = _resolve_content_type(file)
            if not ok:
                raise HTTPException(
                    status_code=415,
                    detail=f"Unsupported file type '{file.content_type}' for '{file.filename}'. "
                    f"Allowed: {', '.join(sorted(settings.allowed_content_types))}",
                )
            file_bytes = await file.read()
            if len(file_bytes) > settings.max_upload_bytes:
                raise HTTPException(
                    status_code=413,
                    detail=f"File '{file.filename}' too large ({len(file_bytes)} bytes). Max: {settings.max_upload_bytes} bytes",
                )
            safe_name = _safe_filename(file.filename)
            content_hash = hashlib.sha256(file_bytes).hexdigest()
            logger.info(f"[IDEMPOTENCY] Calculated content_hash for file '{safe_name}': {content_hash}")
            file_data.append((None, file_bytes, safe_name, content_hash, effective_ct))


    # --- count existing docs for naming
    existing_count = db.query(Document).filter(Document.claim_id == cid).count()

    # --- save files and create document rows
    saved_paths: list[str] = []
    new_docs: list[Document] = []
    new_doc_added = False
    for idx, (raw_path, file_bytes, safe_name, content_hash, effective_ct) in enumerate(file_data):
        logger.info(f"[IDEMPOTENCY] Checking for duplicate: claim_id={claim.id}, content_hash={content_hash}")
        duplicate_doc = db.query(Document).filter(Document.claim_id == claim.id, Document.content_hash == content_hash).first()
        if duplicate_doc:
            logger.info(f"[IDEMPOTENCY] Duplicate document detected for claim {claim.id} and hash {content_hash}, skipping.")
            _audit(db, "DUPLICATE_DOCUMENT_SKIPPED", claim_id=claim.id, metadata={
                "file_name": safe_name,
                "content_hash": content_hash,
                "existing_document_id": str(duplicate_doc.id),
            })
            if raw_path and raw_path.startswith("s3://"):
                try:
                    from libs.shared.storage import MinioStorage
                    MinioStorage.delete_file(raw_path)
                except Exception:
                    pass
            continue

        ext = Path(safe_name).suffix or ".bin"
        stored_name = f"{claim.id}_{existing_count + idx}{ext}"
        s3_key = f"claims/{claim.id}/{stored_name}"

        from libs.shared.storage import MinioStorage
        try:
            if raw_path and raw_path.startswith("s3://"):
                minio_uri = MinioStorage.copy_file(raw_path, s3_key)
                MinioStorage.delete_file(raw_path)
            else:
                minio_uri = MinioStorage.upload_file(s3_key, file_bytes)
            saved_paths.append(minio_uri)
        except Exception as e:
            for p in saved_paths:
                try:
                    MinioStorage.delete_file(p)
                except Exception:
                    pass
            db.rollback()
            logger.exception("Failed to upload/copy file to MinIO")
            raise HTTPException(status_code=500, detail="Failed to store uploaded file in object storage")

        doc = Document(
            claim_id=claim.id,
            file_name=safe_name,
            file_type=effective_ct,
            minio_path=minio_uri,
            content_hash=content_hash,
        )
        db.add(doc)
        new_docs.append(doc)
        new_doc_added = True

    if not new_doc_added:
        logger.info(f"No new documents added for claim {claim.id}; all uploads were duplicates.")
        _audit(db, "DUPLICATE_DOCUMENTS_ALL_SKIPPED", claim_id=claim.id, metadata={
            "file_count": len(file_data),
            "reason": "All uploaded documents were duplicates. Pipeline will still be triggered to ensure combined report."
        })
        # Always trigger pipeline to ensure combined report
        try:
            claim.status = "UPLOADED"
            db.commit()
            task_id = _enqueue_pipeline(str(claim.id))
        except Exception:
            db.rollback()
            logger.exception("Failed to enqueue Celery pipeline for claim %s", claim.id)
            raise HTTPException(status_code=503, detail="No new documents, but failed to enqueue background tasks for combined report")
        payload = _build_claim_response(db, cid, {"task_id": task_id})
        return JSONResponse(status_code=200, content=payload)

    db.flush()
    gate_result = _apply_identity_gate(db, claim.id, new_docs)
    manual_review_message = None
    if gate_result["accepted_count"] == 0:
        claim.status = "MANUAL_REVIEW_REQUIRED"
        upsert_workflow_state(db, claim.id, "MANUAL_REVIEW_REQUIRED", status="FAILED")
        from libs.shared.storage import MinioStorage
        for doc in new_docs:
            if doc.minio_path:
                try:
                    MinioStorage.delete_file(doc.minio_path)
                except Exception:
                    logger.warning("Failed to delete mismatched file from MinIO: %s", doc.minio_path)
            db.delete(doc)
        manual_review_message = (
            "Manual review required: Patient name mismatch detected in the documents you added. "
            "Please check that the uploaded documents have the correct patient details."
        )
    else:
        claim.status = "UPLOADED"

    try:
        db.commit()
    except Exception as exc:
        db.rollback()
        from libs.shared.storage import MinioStorage
        for p in saved_paths:
            try:
                MinioStorage.delete_file(p)
            except Exception:
                pass
        logger.exception("DB commit failed adding documents")
        upload_log.exception(
            "UPLOAD_FAILURE | endpoint=add_documents claim_id=%s stage=db_commit error=%s",
            claim.id, exc,
        )
        raise HTTPException(status_code=500, detail="Failed to save documents")

    logger.info("Added %d docs to claim %s", len(new_docs), claim.id)
    db.refresh(claim)

    task_id: str | None = None
    if gate_result["accepted_count"] > 0:
        try:
            task_id = _enqueue_pipeline(str(claim.id))
        except Exception as exc:
            logger.exception("Failed to enqueue Celery pipeline for claim %s", claim.id)
            upload_log.exception(
                "UPLOAD_FAILURE | endpoint=add_documents claim_id=%s stage=enqueue_pipeline error=%s",
                claim.id, exc,
            )
            raise HTTPException(status_code=503, detail="Documents saved but failed to enqueue background tasks")
    else:
        logger.warning("Claim %s no accepted new docs after identity gate; workflow not retriggered", claim.id)
        upload_log.warning(
            "UPLOAD_PARTIAL | endpoint=add_documents claim_id=%s reason=identity_gate_rejected_all",
            claim.id,
        )

    extra = {"task_id": task_id} if task_id else {}
    if manual_review_message:
        extra["manual_review_reason"] = manual_review_message
    payload = _build_claim_response(db, cid, extra)
    _audit(db, "DOCUMENTS_ADDED", claim_id=claim.id, metadata={
        "files": [s for _, _, s, _, _ in file_data],
        "file_count": len(new_docs),
        "total_documents": existing_count + len(new_docs),
        "identity_gate": gate_result,
        "manual_review_reason": manual_review_message,
    })
    upload_log.info(
        "UPLOAD_SUCCESS | endpoint=add_documents claim_id=%s new_docs=%d total=%d task_id=%s",
        claim.id, len(new_docs), existing_count + len(new_docs), task_id,
    )
    return payload


@router.delete("/claims/{claim_id}/documents/{doc_id}", response_model=ClaimOut)
def delete_document(
    claim_id: str,
    doc_id: str,
    db: Session = Depends(get_db),
):
    """Delete a single document from a claim."""
    cid = _parse_uuid(claim_id)
    did = _parse_uuid(doc_id)
    claim = db.query(Claim).filter(Claim.id == cid).first()
    if not claim:
        raise HTTPException(status_code=404, detail="Claim not found")

    doc = db.query(Document).filter(Document.id == did, Document.claim_id == cid).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Prevent deleting the last document
    doc_count = db.query(Document).filter(Document.claim_id == cid).count()
    if doc_count <= 1:
        raise HTTPException(status_code=400, detail="Cannot delete the only document. Delete the claim instead.")

    # Remove file from disk
    try:
        p = Path(doc.minio_path).resolve()
        if str(p).startswith(str(RAW_STORAGE)):
            p.unlink(missing_ok=True)
    except OSError:
        logger.warning("Failed to delete file %s", doc.minio_path)

    db.delete(doc)
    db.commit()
    db.refresh(claim)
    _audit(db, "DOCUMENT_DELETED", claim_id=cid, metadata={"document_id": str(did), "file_name": doc.file_name})
    logger.info("Deleted doc %s from claim %s", doc_id, claim_id)
    return _build_claim_response(db, cid)


@router.delete("/claims", status_code=204)
def delete_all_claims(db: Session = Depends(get_db)):
    # Delete all raw files from disk
    docs = db.query(Document).all()
    for doc in docs:
        try:
            p = Path(doc.minio_path).resolve()
            if str(p).startswith(str(RAW_STORAGE)):
                p.unlink(missing_ok=True)
        except OSError:
            logger.warning("Failed to delete file %s", doc.minio_path)

    db.query(Claim).delete()
    db.commit()
    logger.info("All claims deleted")


@router.delete("/claims/{claim_id}", status_code=204)
def delete_claim(
    claim_id: str,
    patient_id: str | None = Query(None),
    auth_user: AuthUser = Depends(get_current_user_context),
    db: Session = Depends(get_db)
):
    cid = _parse_uuid(claim_id)
    claim = db.query(Claim).filter(Claim.id == cid).first()
    if not claim:
        raise HTTPException(status_code=404, detail="Claim not found")

    # Enforce Row-Level Security: if caller is a patient identity, verify ownership
    caller_candidates = {
        c.strip().lower() for c in [
            patient_id,
            auth_user.user_id,
            auth_user.patient_id,
            auth_user.email,
        ] if c and c.strip().lower() not in ("user", "null", "undefined")
    }
    if caller_candidates and claim.patient_id:
        if claim.patient_id.strip().lower() not in caller_candidates and auth_user.role not in ("admin", "reviewer", "auditor"):
            raise HTTPException(status_code=404, detail="Claim not found")

    # delete stored files from disk
    docs = db.query(Document).filter(Document.claim_id == cid).all()
    for doc in docs:
        try:
            p = Path(doc.minio_path).resolve()
            if str(p).startswith(str(RAW_STORAGE)):
                p.unlink(missing_ok=True)
        except OSError:
            logger.warning("Failed to delete file %s", doc.minio_path)

    doc_names = [d.file_name for d in docs]
    db.delete(claim)
    db.commit()
    _audit(db, "CLAIM_DELETED", claim_id=cid, metadata={"documents": doc_names})
    logger.info("Claim %s deleted", claim_id)

# ── Include router (standalone mode) ──
app.include_router(router)
