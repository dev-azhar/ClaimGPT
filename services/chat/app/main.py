from __future__ import annotations

import logging
import uuid
from typing import Any

from fastapi import APIRouter, Depends, FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy.sql import func

from .config import settings
from .db import SessionLocal, check_db_health, engine
from .llm import call_llm, get_suggestions, stream_llm
from .models import (
    ChatMessage,
    Claim,
    Document,
    MedicalCode,
    MedicalEntity,
    OcrResult,
    ParsedField,
    Prediction,
    Validation,
)
from .schemas import ChatHistoryOut, ChatMessageOut, ChatRequest, ChatResponse, FieldAction, FieldActionRequest

# ------------------------------------------------------------------ logging
logging.basicConfig(
    level=settings.log_level.upper(),
    format="%(asctime)s  %(levelname)-8s  %(name)s  %(message)s",
)
logger = logging.getLogger("chat")

app = FastAPI(title="ClaimGPT Chat Service")

app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.cors_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ------------------------------------------------------------------ observability
try:
    import os
    import sys
    sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "..", ".."))
    from libs.observability.metrics import PrometheusMiddleware, init_metrics, metrics_endpoint
    from libs.observability.tracing import init_tracing, instrument_fastapi
    init_tracing("chat")
    init_metrics("chat")
    instrument_fastapi(app)
    app.add_middleware(PrometheusMiddleware)
    _metrics_handler = metrics_endpoint()
    if _metrics_handler:
        app.get("/metrics")(_metrics_handler)
except Exception:
    logger.debug("Observability libs not available — skipping")


@app.on_event("shutdown")
def _shutdown():
    engine.dispose()


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


def _parse_uuid(value: str) -> uuid.UUID:
    try:
        return uuid.UUID(value)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid UUID")


# ------------------------------------------------------------------ helpers

def _search_ocr_for_query(
    full_text: str,
    pages: list[dict[str, Any]],
    query: str,
) -> str:
    """
    Question-aware retrieval: search ALL OCR text for content relevant
    to the user's question. Returns the most relevant chunks so the LLM
    can answer precisely about any part of the document.
    """
    import re

    if not full_text:
        return ""

    if not query or not query.strip():
        # No specific question — return first portion + last portion for general context
        if len(full_text) <= 12000:
            return full_text
        return full_text[:8000] + "\n\n[...middle content omitted...]\n\n" + full_text[-4000:]

    query_lower = query.lower()

    # Build query keywords (remove stop words)
    stop_words = {
        "the", "is", "at", "which", "on", "a", "an", "and", "or", "but",
        "in", "with", "to", "for", "of", "this", "that", "what", "how",
        "where", "when", "who", "me", "my", "i", "it", "was", "were",
        "are", "be", "been", "being", "do", "does", "did", "have", "has",
        "had", "can", "could", "will", "would", "shall", "should", "may",
        "might", "about", "from", "into", "show", "tell", "give",
        "please", "find", "get", "look", "see", "any", "all",
    }
    words = re.findall(r"[a-z0-9]+", query_lower)
    keywords = [w for w in words if w not in stop_words and len(w) > 2]

    if not keywords:
        if len(full_text) <= 12000:
            return full_text
        return full_text[:8000] + "\n\n[...]\n\n" + full_text[-4000:]

    # Score each page for relevance
    if pages:
        scored = []
        for p in pages:
            txt_lower = p["text"].lower()
            score = sum(txt_lower.count(kw) for kw in keywords)
            scored.append((score, p["page"], p["text"]))
        scored.sort(key=lambda x: (-x[0], x[1]))

        # Take top-scoring pages + always include first page for context
        result_parts = []
        included_pages = set()
        budget = 12000

        # Always include page 1 for general context
        if pages:
            first = pages[0]
            result_parts.append(f"[Page {first['page']}]\n{first['text']}")
            included_pages.add(first["page"])
            budget -= len(first["text"])

        # Add highest-scoring pages
        for _score, pg, txt in scored:
            if pg in included_pages:
                continue
            if budget <= 0:
                break
            result_parts.append(f"[Page {pg}]\n{txt}")
            included_pages.add(pg)
            budget -= len(txt)

        return "\n\n".join(result_parts)

    # No page structure — search by chunks
    chunk_size = 1000
    overlap = 200
    chunks = []
    for start in range(0, len(full_text), chunk_size - overlap):
        chunk = full_text[start : start + chunk_size]
        if chunk.strip():
            chunks.append((start, chunk))

    scored_chunks = []
    for start, chunk in chunks:
        chunk_lower = chunk.lower()
        score = sum(chunk_lower.count(kw) for kw in keywords)
        scored_chunks.append((score, start, chunk))
    scored_chunks.sort(key=lambda x: -x[0])

    result_parts = []
    budget = 12000
    # Always include the start of the document
    if full_text[:1000].strip():
        result_parts.append(full_text[:1000])
        budget -= 1000

    for score, _start, chunk in scored_chunks:
        if score == 0:
            break
        if budget <= 0:
            break
        result_parts.append(chunk)
        budget -= len(chunk)

    return "\n\n".join(result_parts) if result_parts else full_text[:12000]

def _read_document_text(file_path: str) -> str:
    """Extract text from any supported document type (PDF, DOCX, Excel, images, text)."""
    from pathlib import Path
    p = Path(file_path)
    suffix = p.suffix.lower()

    try:
        # PDF extraction
        if suffix == ".pdf":
            import pdfplumber
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    # Also extract tables
                    try:
                        tables = page.extract_tables()
                        for table in (tables or []):
                            for row in table:
                                cells = [str(c).strip() if c else "" for c in row]
                                text_parts.append(" | ".join(cells))
                    except Exception:
                        pass
            return "\n".join(text_parts)

        # Word documents
        if suffix in (".docx", ".doc"):
            try:
                import docx
                doc = docx.Document(file_path)
                parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        parts.append(para.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        parts.append(" | ".join(cells))
                return "\n".join(parts)
            except ImportError:
                return ""

        # Excel
        if suffix in (".xlsx", ".xls"):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                parts = []
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    parts.append(f"[Sheet: {sheet_name}]")
                    for row in ws.iter_rows(values_only=True):
                        cells = [str(c).strip() if c is not None else "" for c in row]
                        if any(cells):
                            parts.append(" | ".join(cells))
                wb.close()
                return "\n".join(parts)
            except ImportError:
                return ""

        # Images — try OCR via pytesseract
        if suffix in (".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp"):
            try:
                import pytesseract
                from PIL import Image
                img = Image.open(file_path)
                text = pytesseract.image_to_string(img)
                return text
            except ImportError:
                return ""

        # Plain text / CSV / JSON / XML / HTML
        if suffix in (".txt", ".csv", ".json", ".xml", ".html", ".htm", ".md", ".log"):
            for enc in ("utf-8", "latin-1"):
                try:
                    return p.read_text(encoding=enc)
                except (UnicodeDecodeError, ValueError):
                    continue
            return ""

    except Exception:
        logger.debug("Failed to read file %s", file_path, exc_info=True)
    return ""


def _get_claim_context(db: Session, claim_id: uuid.UUID, user_query: str = "") -> dict[str, Any] | None:
    """Build comprehensive claim context with full document text and question-aware retrieval."""
    claim = db.query(Claim).filter(Claim.id == claim_id).first()
    if not claim:
        return None

    pf = db.query(ParsedField).filter(ParsedField.claim_id == claim_id).all()
    # Build fields dict — keep the first (primary) value for duplicate field names
    fields: dict[str, Any] = {}
    for r in pf:
        if r.field_name not in fields:
            fields[r.field_name] = r.field_value

    docs = db.query(Document).filter(Document.claim_id == claim_id).all()
    doc_ids = [d.id for d in docs]

    # ── Fetch ALL OCR pages (no limit) ──
    all_ocr_pages: list[dict[str, Any]] = []
    full_ocr_text = ""
    if doc_ids:
        rows = (
            db.query(OcrResult)
            .filter(OcrResult.document_id.in_(doc_ids))
            .order_by(OcrResult.page_number)
            .all()
        )
        for r in rows:
            if r.text and r.text.strip():
                all_ocr_pages.append({
                    "page": r.page_number or 0,
                    "text": r.text.strip(),
                    "confidence": r.confidence,
                })
        full_ocr_text = "\n\n".join(p["text"] for p in all_ocr_pages)

    # If no OCR text yet, read the raw file directly
    if not full_ocr_text and docs:
        for doc in docs:
            fpath = doc.minio_path
            if fpath:
                full_ocr_text = _read_document_text(fpath)
                if full_ocr_text:
                    break

    # ── Question-aware retrieval: find relevant chunks ──
    relevant_text = _search_ocr_for_query(full_ocr_text, all_ocr_pages, user_query)

    # Predictions
    preds = db.query(Prediction).filter(Prediction.claim_id == claim_id).order_by(Prediction.created_at.desc()).limit(3).all()
    predictions = []
    for p in preds:
        predictions.append({
            "rejection_score": p.rejection_score,
            "top_reasons": p.top_reasons,
            "model_name": p.model_name,
        })

    # Validations
    vals = db.query(Validation).filter(Validation.claim_id == claim_id).all()
    validations = []
    for v in vals:
        validations.append({
            "rule_id": v.rule_id,
            "rule_name": v.rule_name,
            "severity": v.severity,
            "message": v.message,
            "passed": str(v.passed).lower() in ("true", "1", "t"),
        })

    # Medical codes
    mc = db.query(MedicalCode).filter(MedicalCode.claim_id == claim_id).all()
    codes = []
    for c in mc:
        codes.append({
            "code": c.code,
            "code_type": c.code_system,
            "description": c.description,
            "confidence": c.confidence,
        })

    # Medical entities (NER extractions)
    entities = db.query(MedicalEntity).filter(MedicalEntity.claim_id == claim_id).all()
    entity_list = []
    for e in entities:
        entity_list.append({
            "text": e.entity_text,
            "type": e.entity_type,
            "confidence": e.confidence,
        })

    return {
        "claim_id": str(claim_id),
        "status": claim.status,
        "policy_id": claim.policy_id,
        "parsed_fields": fields,
        "full_ocr_text": full_ocr_text,
        "relevant_text": relevant_text,
        "ocr_page_count": len(all_ocr_pages),
        "predictions": predictions,
        "validations": validations,
        "medical_codes": codes,
        "medical_entities": entity_list,
    }


# ------------------------------------------------------------------ routes

router = APIRouter()


@router.get("/health")
def health():
    db_ok = check_db_health()
    return {"status": "ok" if db_ok else "degraded", "database": "up" if db_ok else "down"}


@router.post("/{session_id}/message", response_model=ChatResponse)
def send_message(
    session_id: str,
    body: ChatRequest,
    db: Session = Depends(get_db),
):
    """
    Send a user message and get an assistant response.
    If claim_id is provided, claim context is injected into the LLM prompt.
    """
    # Resolve claim context
    claim_id = None
    claim_context = None
    if body.claim_id:
        claim_id = _parse_uuid(body.claim_id)
        claim_context = _get_claim_context(db, claim_id, user_query=body.message)

    # Save user message
    user_msg = ChatMessage(
        claim_id=claim_id,
        role="USER",
        message=body.message,
    )
    db.add(user_msg)
    db.commit()

    # Gather history for this session (last 20 messages)
    history_rows = (
        db.query(ChatMessage)
        .filter(ChatMessage.claim_id == claim_id)
        .order_by(ChatMessage.created_at.desc())
        .limit(20)
        .all()
    )
    history_rows.reverse()

    messages = [
        {"role": r.role.lower(), "content": r.message}
        for r in history_rows
        if r.message
    ]

    # Call LLM
    assistant_text = call_llm(messages, claim_context)

    # Get contextual follow-up suggestions
    suggestions = get_suggestions(body.message, claim_context)

    # Detect field edit intent (add/modify/delete)
    field_actions = _detect_field_actions(body.message, claim_context)

    # Save assistant response
    assistant_msg = ChatMessage(
        claim_id=claim_id,
        role="ASSISTANT",
        message=assistant_text,
    )
    db.add(assistant_msg)
    db.commit()

    return ChatResponse(
        session_id=session_id,
        role="ASSISTANT",
        message=assistant_text,
        claim_id=body.claim_id,
        suggestions=suggestions,
        field_actions=field_actions,
    )


@router.post("/{session_id}/stream")
async def stream_message(
    session_id: str,
    body: ChatRequest,
    db: Session = Depends(get_db),
):
    """
    Stream an assistant response token-by-token via Server-Sent Events.
    The UI receives each chunk in real-time for a ChatGPT-like experience.
    """
    claim_id = None
    claim_context = None
    if body.claim_id:
        claim_id = _parse_uuid(body.claim_id)
        claim_context = _get_claim_context(db, claim_id, user_query=body.message)

    # Save user message
    user_msg = ChatMessage(claim_id=claim_id, role="USER", message=body.message)
    db.add(user_msg)
    db.commit()

    # Gather history
    history_rows = (
        db.query(ChatMessage)
        .filter(ChatMessage.claim_id == claim_id)
        .order_by(ChatMessage.created_at.desc())
        .limit(20)
        .all()
    )
    history_rows.reverse()
    messages = [
        {"role": r.role.lower(), "content": r.message}
        for r in history_rows
        if r.message
    ]

    collected_chunks: list[str] = []

    async def event_generator():
        async for chunk in stream_llm(messages, claim_context):
            # Collect text chunks (not [DONE]) for saving
            if chunk.startswith("data: ") and "[DONE]" not in chunk:
                import json as _j
                try:
                    data = _j.loads(chunk.split("data: ", 1)[1])
                    collected_chunks.append(data.get("content", ""))
                except Exception:
                    pass
            yield chunk

        # Save the full response after streaming completes
        full_text = "".join(collected_chunks)
        if full_text:
            asst = ChatMessage(claim_id=claim_id, role="ASSISTANT", message=full_text)
            db.add(asst)
            db.commit()

        # Send suggestions as the final event
        suggestions = get_suggestions(body.message, claim_context)
        field_actions = _detect_field_actions(body.message, claim_context)
        import json as _j
        yield f"data: {_j.dumps({'suggestions': suggestions, 'field_actions': [a.model_dump() for a in field_actions]})}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


@router.get("/providers")
def list_providers():
    """Return the LLM provider info."""
    from .config import settings as s
    return {
        "current": "ollama",
        "available": [
            {"id": "ollama", "name": "Ollama (Local)", "model": s.ollama_model, "configured": True},
        ],
    }


@router.get("/{session_id}/history", response_model=ChatHistoryOut)
def get_history(
    session_id: str,
    claim_id: str | None = None,
    db: Session = Depends(get_db),
):
    """Retrieve conversation history."""
    query = db.query(ChatMessage)
    if claim_id:
        cid = _parse_uuid(claim_id)
        query = query.filter(ChatMessage.claim_id == cid)

    rows = query.order_by(ChatMessage.created_at).limit(100).all()

    return ChatHistoryOut(
        session_id=session_id,
        messages=[
            ChatMessageOut(
                id=r.id,
                role=r.role,
                message=r.message,
                created_at=r.created_at,
            )
            for r in rows
        ],
    )


# ------------------------------------------------------------------ field action detection

import re as _re


def _detect_field_actions(
    user_message: str,
    claim_context: dict[str, Any] | None,
) -> list[FieldAction]:
    """
    Detect add/modify/delete intent from user message and return structured actions.
    Supports patterns like:
      - "patient name is John Doe" / "add patient name John Doe"
      - "change diagnosis to Diabetes" / "update hospital to ABC Hospital"
      - "remove policy id" / "delete the provider name"
      - "name is missing, it should be Rahul"
    """
    if not claim_context:
        return []

    msg = user_message.strip()
    msg_lower = msg.lower()
    existing_fields = claim_context.get("parsed_fields", {})
    actions: list[FieldAction] = []

    # Known field aliases → canonical field names
    _FIELD_ALIASES = {
        "patient name": "patient_name", "patient_name": "patient_name", "name": "patient_name",
        "patient age": "patient_age", "age": "patient_age",
        "patient gender": "patient_gender", "gender": "patient_gender", "sex": "patient_gender",
        "date of birth": "date_of_birth", "dob": "date_of_birth",
        "hospital": "hospital_name", "hospital name": "hospital_name",
        "doctor": "doctor_name", "doctor name": "doctor_name", "provider": "doctor_name",
        "provider name": "doctor_name", "physician": "doctor_name",
        "diagnosis": "diagnosis", "primary diagnosis": "diagnosis",
        "admission date": "admission_date", "admitted": "admission_date",
        "discharge date": "discharge_date", "discharged": "discharge_date",
        "policy id": "policy_id", "policy": "policy_id", "policy number": "policy_id",
        "claim amount": "claim_amount", "amount": "claim_amount", "total amount": "claim_amount",
        "bill amount": "bill_amount", "billed amount": "bill_amount",
        "room charges": "room_charges", "room charge": "room_charges",
        "procedure": "procedure", "surgery": "procedure",
        "insurance": "insurance_company", "insurance company": "insurance_company",
        "tpa": "tpa_name", "tpa name": "tpa_name",
        "member id": "member_id", "member": "member_id",
        "address": "address", "patient address": "address",
        "phone": "phone", "contact": "phone", "mobile": "phone",
        "email": "email",
    }

    def _resolve_field(text: str) -> str | None:
        t = text.strip().lower()
        if t in _FIELD_ALIASES:
            return _FIELD_ALIASES[t]
        # Try fuzzy: replace spaces with underscore
        t_under = t.replace(" ", "_")
        if t_under in _FIELD_ALIASES:
            return _FIELD_ALIASES[t_under]
        # Direct match in existing fields
        for k in existing_fields:
            if k.lower() == t or k.lower() == t_under:
                return k
        return t_under  # use as-is

    # ── DELETE patterns ──
    del_patterns = [
        r"(?:please\s+)?(?:remove|delete|clear|erase|drop)\s+(?:the\s+)?(.+?)(?:\s+field)?$",
    ]
    for pat in del_patterns:
        m = _re.search(pat, msg_lower)
        if m:
            field = _resolve_field(m.group(1))
            if field and field in existing_fields:
                actions.append(FieldAction(
                    action="delete",
                    field_name=field,
                    old_value=existing_fields.get(field),
                    new_value=None,
                ))
                return actions

    # ── MODIFY / UPDATE patterns ──
    mod_patterns = [
        r"(?:please\s+)?(?:change|update|modify|set|correct|replace)\s+(?:the\s+)?(.+?)\s+(?:to|as|with|=)\s+(.+)",
        r"(?:the\s+)?(.+?)\s+(?:should be|is actually|is supposed to be|needs to be|must be)\s+(.+)",
        r"(?:please\s+)?(?:fix|correct)\s+(?:the\s+)?(.+?)\s+(?:to|it'?s?)\s+(.+)",
    ]
    for pat in mod_patterns:
        m = _re.search(pat, msg_lower)
        if m:
            field = _resolve_field(m.group(1))
            # Get the new value from the ORIGINAL message to preserve case
            raw_match = _re.search(pat, msg, _re.IGNORECASE)
            new_val = raw_match.group(2).strip().rstrip(".!") if raw_match else m.group(2).strip().rstrip(".!")
            if field:
                old_val = existing_fields.get(field)
                actions.append(FieldAction(
                    action="modify" if old_val else "add",
                    field_name=field,
                    old_value=old_val,
                    new_value=new_val,
                ))
                return actions

    # ── ADD patterns ──
    add_patterns = [
        r"(?:please\s+)?(?:add|enter|put|insert|include)\s+(?:the\s+)?(.+?)\s+(?:as|=|:)\s+(.+)",
        r"(?:please\s+)?(?:add|enter|put|insert)\s+(.+?)\s+(?:for|in)\s+(.+)",
        r"(?:the\s+)?(.+?)\s+is\s+(?:missing|empty|blank|not filled).*?(?:it(?:'?s| is| should be)\s+(.+))",
        r"(.+?)\s+is\s+(.+?)$",  # e.g., "patient name is John Doe"
    ]
    for i, pat in enumerate(add_patterns):
        m = _re.search(pat, msg_lower)
        if m:
            if i == 3:
                # "X is Y" — only if X looks like a field name
                candidate = m.group(1).strip()
                if len(candidate.split()) > 4:
                    continue  # too long to be a field name
                resolved = _resolve_field(candidate)
                if not resolved or resolved not in _FIELD_ALIASES.values():
                    # Only match if it's a known field alias
                    found = False
                    for alias in _FIELD_ALIASES:
                        if alias in candidate.lower():
                            found = True
                            break
                    if not found:
                        continue

            field = _resolve_field(m.group(1))
            raw_match = _re.search(pat, msg, _re.IGNORECASE)
            new_val = raw_match.group(2).strip().rstrip(".!") if raw_match else m.group(2).strip().rstrip(".!")

            if field and new_val:
                old_val = existing_fields.get(field)
                actions.append(FieldAction(
                    action="modify" if old_val else "add",
                    field_name=field,
                    old_value=old_val,
                    new_value=new_val,
                ))
                return actions

    return actions


# ------------------------------------------------------------------ apply field actions

@router.post("/fields/apply")
def apply_field_actions(
    body: FieldActionRequest,
    db: Session = Depends(get_db),
):
    """
    Apply add/modify/delete actions on claim parsed fields.
    Called by the UI when user confirms a field change suggested in chat.
    """
    claim_id = _parse_uuid(body.claim_id)
    claim = db.query(Claim).filter(Claim.id == claim_id).first()
    if not claim:
        raise HTTPException(status_code=404, detail="Claim not found")

    results = []
    for act in body.actions:
        fname = act.field_name.strip()

        if act.action == "delete":
            rows = db.query(ParsedField).filter(
                ParsedField.claim_id == claim_id,
                ParsedField.field_name == fname,
            ).all()
            if rows:
                for r in rows:
                    db.delete(r)
                results.append({"field": fname, "action": "deleted", "ok": True})
            else:
                results.append({"field": fname, "action": "delete", "ok": False, "reason": "not found"})

        elif act.action in ("add", "modify"):
            existing = db.query(ParsedField).filter(
                ParsedField.claim_id == claim_id,
                ParsedField.field_name == fname,
            ).first()

            if existing:
                existing.field_value = act.new_value
                results.append({"field": fname, "action": "updated", "ok": True,
                                "old": act.old_value, "new": act.new_value})
            else:
                new_field = ParsedField(
                    claim_id=claim_id,
                    field_name=fname,
                    field_value=act.new_value,
                )
                db.add(new_field)
                results.append({"field": fname, "action": "added", "ok": True,
                                "value": act.new_value})

    db.commit()

    # Also update claim's updated_at timestamp
    claim.updated_at = func.now()
    db.commit()

    return {"status": "ok", "results": results}


# ── Include router (standalone mode) ──
app.include_router(router)
